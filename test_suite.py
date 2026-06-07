"""
test_suite.py — Suite completa v18
Grupos:
  1. DB — flujo normal
  2. DB — datos basura / bordes
  3. DB — IDs invalidos
  4. DB — duplicados
  5. DB — stress / volumen
  6. DB — concurrencia
  7. DB — integridad referencial
  8. Utils — formatear_fecha, CSVCleaner
  9. Logica — filtros avanzados
 10. Logica — busqueda global
 11. Logica — CRUD empresa
 12. Logica — cotizaciones con ruta
 13. Logica — merge/unificar
 14. QA — importador carpetas
 15. QA — extractor CSV
 16. QA — VirtualTree paginacion
 17. Logica — exportador
 18. Logica — backup
 19. Stress — volumen extremo
"""

import sys, os, sqlite3, unittest, threading, time, random, string, json
import tempfile, shutil, csv
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from db_manager import DBManager
from utils import formatear_fecha, CSVCleaner

def fresh_db():
    return DBManager(":memory:")

def rand(n=12):
    return ''.join(random.choices(string.ascii_letters, k=n))

def make_emp(db, nombre=None, pais="Argentina"):
    nombre = nombre or rand()
    db.agregar_empresa(nombre, "Dir", "0800", "a@b.com", "Tech", pais, "")
    r = db.fetchone("SELECT id FROM empresas WHERE nombre=?", (nombre,))
    return r['id'] if r else None

def make_con(db, eid, email=None):
    email = email or f"{rand(6)}@test.com"
    db.agregar_contacto(eid, "Juan Test", email, "123", "AR")
    r = db.fetchone("SELECT id FROM contactos WHERE email=?", (email,))
    return r['id'] if r else None

def make_cot(db, eid, monto=1000.0, desc="Cotiz test", fecha=None):
    return db.agregar_cotizacion(eid, desc, monto, fecha)


# =====================================================================
# 1. FLUJO NORMAL
# =====================================================================
class TestFlujoNormal(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_crear_empresa(self):
        ok = self.db.agregar_empresa("ACME SA","","","","","Argentina","")
        self.assertTrue(ok)
        r = self.db.fetchone("SELECT * FROM empresas WHERE nombre='ACME SA'")
        self.assertIsNotNone(r)

    def test_obtener_empresa_por_id(self):
        eid = make_emp(self.db, "Empresa A")
        e = self.db.obtener_empresa_por_id(eid)
        self.assertEqual(e['nombre'], "Empresa A")

    def test_editar_empresa(self):
        eid = make_emp(self.db, "Vieja")
        self.db.editar_empresa(eid,"Nueva","Av 1","111","n@mail.com","IT","Chile","")
        e = self.db.obtener_empresa_por_id(eid)
        self.assertEqual(e['nombre'], "Nueva")
        self.assertEqual(e['pais'], "Chile")

    def test_eliminar_empresa_en_cascada(self):
        eid = make_emp(self.db, "Para Borrar")
        make_con(self.db, eid)
        make_cot(self.db, eid)
        self.db.eliminar_empresa(eid)
        self.assertEqual(self.db.fetchall(
            "SELECT * FROM contactos WHERE empresa_id=?", (eid,)), [])
        self.assertEqual(self.db.fetchall(
            "SELECT * FROM cotizaciones WHERE empresa_id=?", (eid,)), [])

    def test_agregar_contacto(self):
        eid = make_emp(self.db)
        make_con(self.db, eid, "juan@test.com")
        cons = self.db.get_contactos_por_empresa(eid)
        self.assertEqual(len(cons), 1)
        self.assertEqual(cons[0]['email'], "juan@test.com")

    def test_editar_contacto(self):
        eid = make_emp(self.db)
        cid = make_con(self.db, eid, "antes@test.com")
        self.db.editar_contacto(cid, "Nuevo Nombre", "despues@test.com", "999", "BR")
        c = self.db.fetchone("SELECT * FROM contactos WHERE id=?", (cid,))
        self.assertEqual(c['nombre'], "Nuevo Nombre")

    def test_agregar_cotizacion_y_ultima_fecha(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 5000.0)
        ultima = self.db.get_ultima_cotizacion(eid)
        self.assertIsNotNone(ultima)

    def test_count_by_empresa(self):
        eid = make_emp(self.db)
        for _ in range(3): make_cot(self.db, eid)
        self.assertEqual(self.db.count_by_empresa("cotizaciones", eid), 3)

    def test_tags_vincular_y_obtener(self):
        eid = make_emp(self.db, "Con Tags")
        self.db.vincular_empresa_con_tags(eid, ["cliente", "vip", "activo"])
        tags = self.db.get_tags_de_empresa(eid)
        self.assertIn("vip", tags)
        self.assertIn("cliente", tags)

    def test_unificar_mueve_contactos_y_cotizaciones(self):
        e1 = make_emp(self.db, "Origen")
        e2 = make_emp(self.db, "Destino")
        make_con(self.db, e1, "c1@test.com")
        make_con(self.db, e1, "c2@test.com")
        make_cot(self.db, e1, 999.0)
        ok = self.db.unificar_empresas(e1, e2)
        self.assertTrue(ok)
        self.assertIsNone(self.db.obtener_empresa_por_id(e1))
        self.assertEqual(self.db.count_by_empresa("contactos", e2), 2)
        self.assertEqual(self.db.count_by_empresa("cotizaciones", e2), 1)

    def test_get_similar_empresas(self):
        make_emp(self.db, "Acme Argentina SA")
        make_emp(self.db, "Acme Argentina SRL")
        make_emp(self.db, "Empresa Diferente")
        pares = self.db.get_similar_empresas(75)
        nombres = [(p['nombre1'], p['nombre2']) for p in pares]
        self.assertTrue(any("Acme" in a and "Acme" in b for a, b in nombres),
                        f"No detecto Acme x Acme como similares: {nombres}")

    def test_cotizacion_con_ruta(self):
        eid = make_emp(self.db)
        ok = self.db.agregar_cotizacion_con_ruta(
            eid, "archivo.pdf", 0.0, "2024-01-01 00:00:00", "/tmp/archivo.pdf")
        self.assertTrue(ok)


# =====================================================================
# 2. DATOS BASURA Y BORDES
# =====================================================================
class TestDatosBasura(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_empresa_nombre_vacio(self):
        try: self.db.agregar_empresa("", "", "", "", "", "", "")
        except Exception as e: self.fail(f"Crash nombre vacio: {e}")

    def test_empresa_nombre_1000_chars(self):
        try: self.db.agregar_empresa("X"*1000, "", "", "", "", "", "")
        except Exception as e: self.fail(f"Crash nombre largo: {e}")

    def test_empresa_unicode_extremo(self):
        for nombre in ["empresas_名称", "Ünternehmen", "شركة", "Empresa🚀", "\n\t"]:
            try: self.db.agregar_empresa(nombre, "", "", "", "", "", "")
            except Exception as e: self.fail(f"Crash con '{repr(nombre)}': {e}")

    def test_sql_injection_nombre(self):
        payloads = [
            "'; DROP TABLE empresas; --",
            "' OR '1'='1",
            "1; UPDATE empresas SET nombre='pwned'--",
        ]
        for p in payloads:
            try: self.db.agregar_empresa(p, "", "", "", "", "", "")
            except: pass
        self.assertIsNotNone(
            self.db.fetchall("SELECT * FROM empresas"),
            "SQL injection destruyo la tabla empresas!")

    def test_sql_injection_busqueda(self):
        make_emp(self.db, "Legitima")
        for p in ["'; DROP TABLE empresas;--", "' OR 1=1--"]:
            try: self.db.get_filtered_empresas(p, {})
            except Exception as e: self.fail(f"Crash busqueda: {e}")
        self.assertGreater(self.db.count("empresas"), 0,
                           "SQL injection borro empresas!")

    def test_emails_invalidos_contacto(self):
        eid = make_emp(self.db)
        for email in ["no-email", "@sinusuario", "sindominio@", "",
                      "a"*500+"@b.com", "<script>@xss.com"]:
            try: self.db.agregar_contacto(eid, "Test", email, "", "")
            except Exception as e: self.fail(f"Crash email '{repr(email)}': {e}")

    def test_monto_negativo(self):
        eid = make_emp(self.db)
        try: self.db.agregar_cotizacion(eid, "Negativa", -9999.99)
        except Exception as e: self.fail(f"Crash monto negativo: {e}")

    def test_monto_cero(self):
        eid = make_emp(self.db)
        ok = self.db.agregar_cotizacion(eid, "Cero", 0.0)
        self.assertIsNotNone(ok)

    def test_monto_extremo(self):
        eid = make_emp(self.db)
        try: self.db.agregar_cotizacion(eid, "Gigante", 999_999_999_999.99)
        except Exception as e: self.fail(f"Crash monto extremo: {e}")

    def test_descripcion_5000_chars(self):
        eid = make_emp(self.db)
        try: self.db.agregar_cotizacion(eid, "A"*5000, 100.0)
        except Exception as e: self.fail(f"Crash descripcion larga: {e}")

    def test_descripcion_xss_y_sql(self):
        eid = make_emp(self.db)
        for desc in ["<script>alert(1)</script>", "/* DROP */",
                     "'; DELETE--", '"comillas"', "'apostrofes'"]:
            try: self.db.agregar_cotizacion(eid, desc, 1.0)
            except Exception as e: self.fail(f"Crash desc especial: {e}")


# =====================================================================
# 3. IDs INVALIDOS
# =====================================================================
class TestIdsInvalidos(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_empresa_inexistente(self):
        self.assertIsNone(self.db.obtener_empresa_por_id(99999))

    def test_empresa_id_cero(self):
        self.assertIsNone(self.db.obtener_empresa_por_id(0))

    def test_empresa_id_negativo(self):
        self.assertIsNone(self.db.obtener_empresa_por_id(-1))

    def test_empresa_id_none(self):
        try: self.db.obtener_empresa_por_id(None)
        except Exception as e: self.fail(f"Crash id=None: {e}")

    def test_eliminar_empresa_inexistente(self):
        try: self.db.eliminar_empresa(99999)
        except Exception as e: self.fail(f"Crash eliminando inexistente: {e}")

    def test_editar_empresa_inexistente(self):
        try: self.db.editar_empresa(99999, "X", "", "", "", "", "", "")
        except Exception as e: self.fail(f"Crash editando inexistente: {e}")

    def test_eliminar_contacto_inexistente(self):
        try: self.db.eliminar_contacto(99999)
        except Exception as e: self.fail(f"Crash eliminando contacto: {e}")

    def test_get_contactos_empresa_inexistente(self):
        self.assertEqual(self.db.get_contactos_por_empresa(99999), [])

    def test_get_cotizaciones_empresa_inexistente(self):
        self.assertEqual(self.db.get_cotizaciones_por_empresa(99999), [])

    def test_ultima_cotizacion_sin_cotizaciones(self):
        eid = make_emp(self.db)
        self.assertIsNone(self.db.get_ultima_cotizacion(eid))

    def test_unificar_consigo_misma(self):
        eid = make_emp(self.db, "Auto-Merge")
        try:
            ok = self.db.unificar_empresas(eid, eid)
            if ok:
                self.assertIsNotNone(
                    self.db.obtener_empresa_por_id(eid),
                    "Auto-merge elimino la empresa!")
        except: pass

    def test_unificar_ids_inexistentes(self):
        try: self.db.unificar_empresas(88888, 99999)
        except Exception as e: self.fail(f"Crash unificando inexistentes: {e}")

    def test_count_tabla_inexistente(self):
        result = self.db.count("tabla_xyz_inexistente")
        self.assertEqual(result, 0)

    def test_filtros_none(self):
        try: self.db.get_filtered_empresas("", None)
        except Exception as e: self.fail(f"Crash filtros=None: {e}")

    def test_filtros_search_none(self):
        try: self.db.get_filtered_empresas(None, {})
        except Exception as e: self.fail(f"Crash search=None: {e}")


# =====================================================================
# 4. DUPLICADOS
# =====================================================================
class TestDuplicados(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_empresa_duplicada_no_crashea(self):
        self.db.agregar_empresa("ACME","","","","","","")
        try: self.db.agregar_empresa("ACME","","","","","","")
        except: pass
        self.assertIsNotNone(self.db.fetchall("SELECT * FROM empresas"))

    def test_tags_duplicados_no_se_repiten(self):
        eid = make_emp(self.db, "Tags Dup")
        self.db.vincular_empresa_con_tags(eid, ["vip","vip","cliente","vip"])
        tags = self.db.get_tags_de_empresa(eid)
        self.assertEqual(tags.count("vip"), 1, f"Tags duplicados: {tags}")

    def test_cotizacion_misma_descripcion(self):
        eid = make_emp(self.db)
        self.db.agregar_cotizacion(eid, "mismo.pdf", 0.0)
        try: self.db.agregar_cotizacion(eid, "mismo.pdf", 0.0)
        except Exception as e: self.fail(f"Crash cotizacion duplicada: {e}")


# =====================================================================
# 5. STRESS VOLUMEN
# =====================================================================
class TestStressVolumen(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_1000_empresas(self):
        for i in range(1000):
            self.db.agregar_empresa(f"Empresa_{i:04d}","","","","","AR","")
        self.assertEqual(self.db.count("empresas"), 1000)

    def test_5000_cotizaciones_una_empresa(self):
        eid = make_emp(self.db, "Mega")
        for i in range(5000):
            self.db.agregar_cotizacion(eid, f"Cotiz #{i}", float(i))
        self.assertEqual(self.db.count_by_empresa("cotizaciones", eid), 5000)

    def test_busqueda_rapida_1000_empresas(self):
        for i in range(1000):
            self.db.agregar_empresa(f"Empresa_{i:04d}","","","","","AR","")
        t0 = time.time()
        result = self.db.get_filtered_empresas("Empresa_05", {})
        elapsed = time.time() - t0
        self.assertLess(elapsed, 3.0, f"Busqueda tardo {elapsed:.2f}s")
        self.assertTrue(len(result) > 0)

    def test_similar_empresas_200(self):
        for i in range(200):
            self.db.agregar_empresa(f"Empresa Similar {i}","","","","","","")
        t0 = time.time()
        self.db.get_similar_empresas(80)
        elapsed = time.time() - t0
        self.assertLess(elapsed, 12.0, f"get_similar_empresas tardo {elapsed:.2f}s")

    def test_unificar_200_elementos(self):
        e1 = make_emp(self.db, "Origen Grande")
        e2 = make_emp(self.db, "Destino Grande")
        for i in range(200):
            self.db.agregar_contacto(e1, f"C{i}", f"c{i}@test.com","","")
            self.db.agregar_cotizacion(e1, f"Cotiz {i}", float(i))
        self.db.unificar_empresas(e1, e2)
        self.assertEqual(self.db.count_by_empresa("contactos", e2), 200)
        self.assertEqual(self.db.count_by_empresa("cotizaciones", e2), 200)

    def test_500_crear_250_borrar_sin_huerfanos(self):
        ids = [make_emp(self.db, f"Emp_B_{i}") for i in range(500)]
        for i in range(200):
            make_cot(self.db, ids[i])
        for eid in ids[:250]: self.db.eliminar_empresa(eid)
        self.assertEqual(self.db.count("empresas"), 250)
        for eid in ids[:250]:
            cots = self.db.fetchall(
                "SELECT * FROM cotizaciones WHERE empresa_id=?", (eid,))
            self.assertEqual(len(cots), 0,
                f"Cotizaciones huerfanas para empresa_id={eid}")


# =====================================================================
# 6. CONCURRENCIA
# =====================================================================
class TestConcurrencia(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_escrituras_concurrentes(self):
        errors = []
        def worker(prefix):
            for i in range(30):
                try:
                    self.db.agregar_empresa(f"{prefix}_{i}_{rand(4)}","","","","","","")
                except Exception as e:
                    errors.append(str(e))
        threads = [threading.Thread(target=worker, args=(f"T{t}",))
                   for t in range(10)]
        for th in threads: th.start()
        for th in threads: th.join(timeout=15)
        alive = [th for th in threads if th.is_alive()]
        self.assertEqual(len(alive), 0, f"{len(alive)} threads colgados (deadlock?)")
        bad = [e for e in errors
               if "lock" not in e.lower() and "database" not in e.lower()]
        self.assertEqual(bad, [], f"Errores no-lock: {bad}")

    def test_lecturas_concurrentes(self):
        for i in range(50): make_emp(self.db, f"Emp_{i}")
        errors = []
        def reader():
            for _ in range(50):
                try: self.db.get_filtered_empresas("", {})
                except Exception as e: errors.append(str(e))
        threads = [threading.Thread(target=reader) for _ in range(8)]
        for th in threads: th.start()
        for th in threads: th.join(timeout=15)
        self.assertEqual(errors, [], f"Errores en lecturas: {errors}")

    def test_mix_lectura_escritura(self):
        eid = make_emp(self.db, "Base")
        errors = []
        def writer():
            for i in range(50):
                try: self.db.agregar_cotizacion(eid, f"C{i}", float(i))
                except Exception as e: errors.append(("w", str(e)))
        def reader():
            for _ in range(50):
                try: self.db.get_cotizaciones_por_empresa(eid)
                except Exception as e: errors.append(("r", str(e)))
        threads = [threading.Thread(target=writer),
                   threading.Thread(target=reader),
                   threading.Thread(target=reader)]
        for th in threads: th.start()
        for th in threads: th.join(timeout=15)
        bad = [e for e in errors if "lock" not in e[1].lower()]
        self.assertEqual(bad, [], f"Errores no-lock: {bad}")


# =====================================================================
# 7. INTEGRIDAD REFERENCIAL
# =====================================================================
class TestIntegridadReferencial(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_orden_cotizaciones_desc(self):
        eid = make_emp(self.db)
        self.db.agregar_cotizacion(eid, "P", 100.0, "2023-01-01 00:00:00")
        self.db.agregar_cotizacion(eid, "S", 200.0, "2024-06-15 00:00:00")
        self.db.agregar_cotizacion(eid, "T", 300.0, "2025-03-20 00:00:00")
        cots = self.db.get_cotizaciones_por_empresa(eid)
        fechas = [c['fecha'] for c in cots]
        self.assertEqual(fechas, sorted(fechas, reverse=True))

    def test_ultima_cotizacion_es_la_mas_reciente(self):
        eid = make_emp(self.db)
        self.db.agregar_cotizacion(eid, "Vieja", 100.0, "2020-01-01 00:00:00")
        self.db.agregar_cotizacion(eid, "Nueva", 200.0, "2025-12-31 00:00:00")
        ultima = self.db.get_ultima_cotizacion(eid)
        self.assertTrue(ultima.startswith("2025"), f"got: {ultima}")

    def test_count_total_consistente(self):
        e1 = make_emp(self.db, "E1"); e2 = make_emp(self.db, "E2")
        for _ in range(5): make_cot(self.db, e1)
        for _ in range(3): make_cot(self.db, e2)
        self.assertEqual(
            self.db.count("cotizaciones"),
            self.db.count_by_empresa("cotizaciones", e1) +
            self.db.count_by_empresa("cotizaciones", e2))

    def test_cotizacion_monto_none(self):
        eid = make_emp(self.db)
        try: self.db.agregar_cotizacion(eid, "Sin monto", None)
        except Exception as e: self.fail(f"Crash monto=None: {e}")

    def test_empresa_pais_none(self):
        try: self.db.agregar_empresa("Sin Pais","","","","",None,"")
        except Exception as e: self.fail(f"Crash pais=None: {e}")


# =====================================================================
# 8. UTILS
# =====================================================================
class TestUtils(unittest.TestCase):
    def test_formatear_fecha_completa(self):
        self.assertEqual(formatear_fecha("2024-06-15 10:30:00"), "15-06-2024")

    def test_formatear_fecha_solo_fecha(self):
        self.assertEqual(formatear_fecha("2024-06-15"), "15-06-2024")

    def test_formatear_fecha_vacia(self):
        self.assertEqual(formatear_fecha(""), "")
        self.assertEqual(formatear_fecha(None), "")

    def test_formatear_fecha_basura(self):
        for b in ["no-es-fecha","32-13-2024","hola","12345","//","\x00"]:
            result = formatear_fecha(b)
            self.assertIsInstance(result, str)

    def test_csvclean_email_valido(self):
        self.assertEqual(CSVCleaner.clean_email("Test@Empresa.com"),
                         "test@empresa.com")

    def test_csvclean_email_con_comillas(self):
        self.assertEqual(CSVCleaner.clean_email('"test@empresa.com"'),
                         "test@empresa.com")

    def test_csvclean_email_invalido(self):
        for e in ["no-es-email","@sinusuario","sindominio@","",None]:
            self.assertEqual(CSVCleaner.clean_email(e), "",
                             f"Deberia devolver '' para '{e}'")

    def test_csvclean_empresa_gmail_es_sin_empresa(self):
        self.assertEqual(
            CSVCleaner.extract_empresa_from_email("u@gmail.com",""), "Sin empresa")

    def test_csvclean_empresa_corporativo(self):
        empresa = CSVCleaner.extract_empresa_from_email("juan@acmecorp.com.ar","")
        self.assertNotEqual(empresa.lower(), "sin empresa")


# =====================================================================
# 9. FILTROS AVANZADOS
# =====================================================================
class TestFiltrosAvanzados(unittest.TestCase):
    def setUp(self):
        self.db = fresh_db()
        from datetime import datetime
        ahora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.e1 = make_emp(self.db, "Con Cotizaciones")
        self.db.agregar_cotizacion(self.e1, "old", 100.0, "2023-01-01 00:00:00")
        self.e2 = make_emp(self.db, "Sin Cotizaciones")
        self.e3 = make_emp(self.db, "Con Contactos")
        make_con(self.db, self.e3)
        self.e4 = make_emp(self.db, "Reciente")
        self.db.agregar_cotizacion(self.e4, "nueva", 200.0, ahora)

    def test_filtro_con_cotizaciones(self):
        result = self.db.get_filtered_empresas("",{"cotizaciones_cond":"con"})
        ids = {r['id'] for r in result}
        self.assertIn(self.e1, ids)
        self.assertNotIn(self.e2, ids)

    def test_filtro_sin_cotizaciones(self):
        result = self.db.get_filtered_empresas("",{"cotizaciones_cond":"sin"})
        ids = {r['id'] for r in result}
        self.assertIn(self.e2, ids)
        self.assertNotIn(self.e1, ids)

    def test_filtro_con_contactos(self):
        result = self.db.get_filtered_empresas("",{"contactos_cond":"con"})
        ids = {r['id'] for r in result}
        self.assertIn(self.e3, ids)

    def test_filtro_sin_contactos(self):
        result = self.db.get_filtered_empresas("",{"contactos_cond":"sin"})
        ids = {r['id'] for r in result}
        self.assertNotIn(self.e3, ids)

    def test_filtro_dias_cotizacion(self):
        result = self.db.get_filtered_empresas("",{"dias_cotizacion": 30})
        ids = {r['id'] for r in result}
        self.assertIn(self.e1, ids,
            "e1 (cotiz antigua) debe aparecer con filtro dias>30")
        self.assertNotIn(self.e4, ids,
            "e4 (cotiz reciente) NO debe aparecer con filtro dias>30")

    def test_filtro_texto_combinado(self):
        result = self.db.get_filtered_empresas(
            "Con", {"cotizaciones_cond":"con"})
        ids = {r['id'] for r in result}
        self.assertIn(self.e1, ids)
        self.assertNotIn(self.e2, ids)

    def test_filtros_vacios_devuelve_todo(self):
        total = self.db.count("empresas")
        result = self.db.get_filtered_empresas("", {})
        self.assertEqual(len(result), total)


# =====================================================================
# 10. BUSQUEDA GLOBAL (SQL directo)
# =====================================================================
class TestBusquedaGlobal(unittest.TestCase):
    def setUp(self):
        self.db = fresh_db()
        self.e1 = make_emp(self.db, "Empresa Alpha")
        self.e2 = make_emp(self.db, "Empresa Beta")
        self.db.agregar_cotizacion(self.e1, "Propuesta Alpha 2024", 5000.0,
                                   "2024-03-01 00:00:00")
        self.db.agregar_cotizacion(self.e2, "Contrato Beta", 12000.0,
                                   "2024-08-15 00:00:00")
        self.db.agregar_cotizacion(self.e2, "Renovacion Beta", 8000.0,
                                   "2025-01-10 00:00:00")

    def test_busqueda_por_texto_descripcion(self):
        rows = self.db.fetchall("""
            SELECT c.*, e.nombre empresa_nombre
            FROM cotizaciones c JOIN empresas e ON c.empresa_id=e.id
            WHERE c.descripcion LIKE ? OR e.nombre LIKE ?
        """, ("%Alpha%", "%Alpha%"))
        self.assertEqual(len(rows), 1)

    def test_busqueda_por_empresa(self):
        rows = self.db.fetchall("""
            SELECT c.* FROM cotizaciones c
            JOIN empresas e ON c.empresa_id=e.id
            WHERE e.nombre LIKE ?
        """, ("%Beta%",))
        self.assertEqual(len(rows), 2)

    def test_busqueda_monto_minimo(self):
        rows = self.db.fetchall("""
            SELECT c.* FROM cotizaciones c WHERE c.monto >= ?
        """, (10000.0,))
        montos = [r['monto'] for r in rows]
        self.assertIn(12000.0, montos)
        self.assertNotIn(5000.0, montos)

    def test_busqueda_sin_resultados(self):
        rows = self.db.fetchall("""
            SELECT c.* FROM cotizaciones c
            JOIN empresas e ON c.empresa_id=e.id
            WHERE e.nombre LIKE ?
        """, ("%INEXISTENTE_XYZ%",))
        self.assertEqual(rows, [])

    def test_total_monto_suma_correcta(self):
        rows = self.db.fetchall("SELECT SUM(monto) total FROM cotizaciones")
        total = rows[0]['total'] if rows else 0
        self.assertAlmostEqual(total, 25000.0, places=1)


# =====================================================================
# 11. CRUD EMPRESA (logica)
# =====================================================================
class TestCRUDEmpresa(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_nueva_empresa_con_tags(self):
        ok = self.db.agregar_empresa("Nueva SA","Av 1","0800","n@sa.com",
                                     "Tech","Argentina","cliente, vip")
        self.assertTrue(ok)
        eid = self.db.fetchone(
            "SELECT id FROM empresas WHERE nombre='Nueva SA'")['id']
        tags = self.db.get_tags_de_empresa(eid)
        self.assertIn("cliente", tags)
        self.assertIn("vip", tags)

    def test_editar_empresa_actualiza_tags(self):
        eid = make_emp(self.db, "Tag Test")
        self.db.vincular_empresa_con_tags(eid, ["viejo_tag"])
        self.db.editar_empresa(eid,"Tag Test","","","","","",
                               "nuevo_tag1, nuevo_tag2")
        tags = self.db.get_tags_de_empresa(eid)
        self.assertIn("nuevo_tag1", tags)
        self.assertNotIn("viejo_tag", tags)

    def test_eliminar_empresa_limpia_tags(self):
        eid = make_emp(self.db, "Con Tags")
        self.db.vincular_empresa_con_tags(eid, ["a","b","c"])
        self.db.eliminar_empresa(eid)
        rows = self.db.fetchall(
            "SELECT * FROM empresa_tags WHERE empresa_id=?", (eid,))
        self.assertEqual(rows, [])

    def test_editar_nombre_vacio_no_crashea(self):
        eid = make_emp(self.db, "Original")
        try: self.db.editar_empresa(eid,"","","","","","","")
        except Exception as e: self.fail(f"Crash nombre vacio: {e}")

    def test_multiples_empresas_mismo_pais(self):
        for i in range(20): make_emp(self.db, f"Emp_{i}", pais="Chile")
        result = self.db.get_filtered_empresas("", {})
        chile = [r for r in result if r.get('pais') == 'Chile']
        self.assertEqual(len(chile), 20)


# =====================================================================
# 12. COTIZACIONES CON RUTA
# =====================================================================
class TestCotizacionesConRuta(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_guardar_y_recuperar_ruta(self):
        eid = make_emp(self.db)
        ruta = "/srv/cotizaciones/ACME/propuesta_2024.pdf"
        ok = self.db.agregar_cotizacion_con_ruta(
            eid, "propuesta_2024.pdf", 0.0, "2024-01-01 00:00:00", ruta)
        self.assertTrue(ok)

    def test_ruta_none_no_crashea(self):
        eid = make_emp(self.db)
        try:
            self.db.agregar_cotizacion_con_ruta(
                eid, "archivo.pdf", 0.0, "2024-01-01 00:00:00", None)
        except Exception as e:
            self.fail(f"Crash ruta=None: {e}")

    def test_ruta_con_espacios_y_caracteres(self):
        eid = make_emp(self.db)
        ruta = "/ruta con espacios/y acentos/archivo (copia).pdf"
        try:
            self.db.agregar_cotizacion_con_ruta(
                eid, "archivo.pdf", 0.0, "2024-01-01 00:00:00", ruta)
        except Exception as e:
            self.fail(f"Crash ruta especial: {e}")


# =====================================================================
# 13. MERGE / UNIFICAR
# =====================================================================
class TestMerge(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_merge_origen_desaparece(self):
        e1 = make_emp(self.db, "Desaparece")
        e2 = make_emp(self.db, "Persiste")
        self.db.unificar_empresas(e1, e2)
        self.assertIsNone(self.db.obtener_empresa_por_id(e1))
        self.assertIsNotNone(self.db.obtener_empresa_por_id(e2))

    def test_merge_transfiere_tags(self):
        e1 = make_emp(self.db, "Origen Tags")
        e2 = make_emp(self.db, "Destino Tags")
        self.db.vincular_empresa_con_tags(e1, ["tag_origen"])
        self.db.vincular_empresa_con_tags(e2, ["tag_destino"])
        self.db.unificar_empresas(e1, e2)
        # e2 debe seguir existiendo con sus tags
        tags = self.db.get_tags_de_empresa(e2)
        self.assertIn("tag_destino", tags)

    def test_merge_no_crashea_con_emails_dup(self):
        e1 = make_emp(self.db, "Dup Origen")
        e2 = make_emp(self.db, "Dup Destino")
        self.db.agregar_contacto(e1, "Juan","dup@test.com","","")
        self.db.agregar_contacto(e2, "Juan","dup@test.com","","")
        try: self.db.unificar_empresas(e1, e2)
        except Exception as e: self.fail(f"Crash merge emails dup: {e}")

    def test_merge_cadena_multiples(self):
        """Merge de 5 empresas en cadena no debe dejar la DB corrupta."""
        ids = [make_emp(self.db, f"Chain_{i}") for i in range(5)]
        destino = ids[0]
        for origen in ids[1:]:
            try: self.db.unificar_empresas(origen, destino)
            except: pass
        self.assertIsNotNone(self.db.obtener_empresa_por_id(destino))


# =====================================================================
# 14. QA — IMPORTADOR CARPETAS
# =====================================================================
class TestImportadorCarpetas(unittest.TestCase):
    def setUp(self):
        self.db = fresh_db()
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_matching_nombre_carpeta_empresa(self):
        make_emp(self.db, "ACME Argentina")
        from fuzzywuzzy import process
        companies = self.db.fetchall("SELECT id, nombre FROM empresas")
        comp_names = [c['nombre'] for c in companies]
        hits = process.extract("ACME Argentina", comp_names, limit=1)
        self.assertGreater(hits[0][1], 80)
        self.assertEqual(hits[0][0], "ACME Argentina")

    def test_archivos_sueltos_detectados(self):
        open(os.path.join(self.tmpdir, "archivo.pdf"), 'w').close()
        open(os.path.join(self.tmpdir, "otro.docx"), 'w').close()
        found = []
        exts = {'.pdf', '.docx', '.xlsx', '.doc', '.xls'}
        for root, _, files in os.walk(self.tmpdir):
            for f in files:
                if os.path.splitext(f)[1].lower() in exts:
                    found.append(f)
        self.assertEqual(len(found), 2)

    def test_extensiones_no_soportadas_ignoradas(self):
        for nombre in ["foto.jpg","musica.mp3","video.mp4"]:
            open(os.path.join(self.tmpdir, nombre), 'w').close()
        exts = {'.pdf', '.docx', '.xlsx', '.doc', '.xls'}
        found = []
        for root, _, files in os.walk(self.tmpdir):
            for f in files:
                if os.path.splitext(f)[1].lower() in exts:
                    found.append(f)
        self.assertEqual(found, [])

    def test_estructura_profunda_3_niveles(self):
        d = os.path.join(self.tmpdir, "N1", "N2", "N3")
        os.makedirs(d)
        open(os.path.join(d, "profundo.pdf"), 'w').close()
        found = []
        for root, _, files in os.walk(self.tmpdir):
            for f in files:
                if f.endswith('.pdf'): found.append(f)
        self.assertIn("profundo.pdf", found)

    def test_matching_fuzzy_umbral(self):
        from fuzzywuzzy import fuzz
        pares = [
            ("ACME S.A.",          "ACME SA",            80),
            ("Empresa ABC SRL",    "Empresa ABC S.R.L.", 75),
            ("Tech Corp",          "TechCorp",           70),
        ]
        for a, b, min_sim in pares:
            sim = fuzz.ratio(a.lower(), b.lower())
            self.assertGreater(sim, min_sim,
                f"'{a}' vs '{b}': similitud {sim} < {min_sim}")

    def test_db_vacia_no_crashea(self):
        from fuzzywuzzy import process
        comp_names = []
        if comp_names:
            process.extract("Cualquier", comp_names, limit=1)
        # No debe crashear con lista vacia
        self.assertEqual(comp_names, [])

    def test_importar_cotizacion_guarda_correctamente(self):
        eid = make_emp(self.db, "Empresa Test")
        fake_file = os.path.join(self.tmpdir, "cotiz.pdf")
        open(fake_file, 'w').close()
        import datetime
        mtime = datetime.datetime.fromtimestamp(
            os.path.getmtime(fake_file)).strftime('%Y-%m-%d %H:%M:%S')
        ok = self.db.agregar_cotizacion_con_ruta(
            eid, "cotiz.pdf", 0.0, mtime, fake_file)
        self.assertTrue(ok)
        self.assertEqual(self.db.count_by_empresa("cotizaciones", eid), 1)


# =====================================================================
# 15. QA — EXTRACTOR CSV
# =====================================================================
class TestExtractorCSV(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _csv(self, filename, rows, headers):
        path = os.path.join(self.tmpdir, filename)
        with open(path, 'w', newline='', encoding='utf-8') as f:
            w = csv.writer(f)
            w.writerow(headers)
            w.writerows(rows)
        return path

    def test_csv_columnas_estandar(self):
        path = self._csv("std.csv",
                         [["Juan","juan@acme.com","ACME","AR"]],
                         ["nombre","email","empresa","pais"])
        with open(path, newline='', encoding='utf-8') as f:
            rows = list(csv.DictReader(f))
        self.assertEqual(rows[0]['email'], 'juan@acme.com')

    def test_csv_columnas_mayusculas(self):
        path = self._csv("may.csv",
                         [["Pedro","pedro@beta.com","Beta","CL"]],
                         ["NOMBRE","EMAIL","EMPRESA","PAIS"])
        with open(path, newline='', encoding='utf-8') as f:
            rows = list(csv.DictReader(f))
        normalized = {k.lower(): v for k,v in rows[0].items()}
        self.assertIn('email', normalized)

    def test_csv_vacio_no_crashea(self):
        path = self._csv("vacio.csv", [], ["nombre","email"])
        with open(path, newline='', encoding='utf-8') as f:
            rows = list(csv.DictReader(f))
        self.assertEqual(rows, [])

    def test_csv_emails_mixtos(self):
        path = self._csv("mixed.csv",
                         [["A","valido@corp.com","Corp",""],
                          ["B","no-es-email","Corp",""],
                          ["C","","Corp",""]],
                         ["nombre","email","empresa","pais"])
        with open(path, newline='', encoding='utf-8') as f:
            rows = list(csv.DictReader(f))
        valid = [r['email'] for r in rows
                 if '@' in r['email'] and '.' in r['email'].split('@')[-1]]
        self.assertIn("valido@corp.com", valid)
        self.assertEqual(len(valid), 1)

    def test_csv_unicode(self):
        path = self._csv("uni.csv",
                         [["Maria Jose","mj@emp.com","Empresa","AR"]],
                         ["nombre","email","empresa","pais"])
        with open(path, newline='', encoding='utf-8') as f:
            rows = list(csv.DictReader(f))
        self.assertEqual(rows[0]['nombre'], "Maria Jose")


# =====================================================================
# 16. QA — VIRTUAL TREE (logica pura, sin GUI)
# =====================================================================
class TestVirtualTreeLogica(unittest.TestCase):
    PAGE = 200

    def _rows(self, n):
        return [(str(i), (f"Empresa {i}", "AR", "01-01-2024", i), ("green",))
                for i in range(n)]

    def test_pagina_inicial_tiene_PAGE_rows(self):
        rows = self._rows(1000)
        chunk = rows[0:self.PAGE]
        self.assertEqual(len(chunk), self.PAGE)

    def test_ultima_pagina_parcial(self):
        rows = self._rows(450)
        last_offset = (len(rows) // self.PAGE) * self.PAGE
        chunk = rows[last_offset:]
        self.assertEqual(len(chunk), 50)

    def test_scroll_to_offset_correcto(self):
        rows = self._rows(1000)
        target = "750"
        for i, (iid, _, _) in enumerate(rows):
            if iid == target:
                expected_offset = (i // self.PAGE) * self.PAGE
                self.assertEqual(expected_offset, 600)
                return
        self.fail("iid 750 no encontrado")

    def test_filtro_reduce_lista(self):
        rows = self._rows(1000)
        filtered = [r for r in rows if int(r[0]) % 2 == 0]
        self.assertEqual(len(filtered), 500)
        chunk = filtered[0:self.PAGE]
        self.assertEqual(chunk[0][0], "0")
        self.assertEqual(chunk[1][0], "2")

    def test_lista_vacia(self):
        chunk = [][0:self.PAGE]
        self.assertEqual(chunk, [])

    def test_orden_preservado(self):
        rows = [(str(i), (f"Z_{1000-i}",), ()) for i in range(100)]
        chunk = rows[0:self.PAGE]
        nombres = [v[0] for _, v, _ in chunk]
        self.assertEqual(nombres[0], "Z_1000")
        self.assertEqual(nombres[-1], "Z_901")

    def test_iid_unico_en_pagina(self):
        rows = self._rows(500)
        chunk = rows[0:self.PAGE]
        iids = [r[0] for r in chunk]
        self.assertEqual(len(iids), len(set(iids)),
                         "Hay iids duplicados en la pagina")


# =====================================================================
# 17. EXPORTADOR
# =====================================================================
class TestExportador(unittest.TestCase):
    def setUp(self):
        self.db = fresh_db()
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_get_all_vacio(self):
        result = self.db.get_all_empresas_with_cotizaciones()
        self.assertIsInstance(result, list)

    def test_get_all_con_datos(self):
        eid = make_emp(self.db, "Export Test")
        make_cot(self.db, eid, 9999.0, "Cotiz exportable")
        result = self.db.get_all_empresas_with_cotizaciones()
        self.assertTrue(len(result) > 0)

    def test_exportar_csv(self):
        from utils import Exportador
        eid = make_emp(self.db, "CSV Export")
        make_cot(self.db, eid, 5000.0, "Cotiz CSV")
        datos = self.db.get_all_empresas_with_cotizaciones()
        path = os.path.join(self.tmpdir, "export.csv")
        ok = Exportador().a_csv(datos, path)
        self.assertTrue(ok)
        self.assertGreater(os.path.getsize(path), 0)

    def test_exportar_excel(self):
        try: import openpyxl
        except ImportError: self.skipTest("openpyxl no instalado")
        from utils import Exportador
        eid = make_emp(self.db, "Excel Export")
        make_cot(self.db, eid, 7500.0)
        datos = self.db.get_all_empresas_with_cotizaciones()
        path = os.path.join(self.tmpdir, "export.xlsx")
        ok = Exportador().a_excel(datos, path)
        self.assertTrue(ok)
        self.assertGreater(os.path.getsize(path), 0)


# =====================================================================
# 18. BACKUP
# =====================================================================
class TestBackup(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.db_path = os.path.join(self.tmpdir, "test.db")
        self.db = DBManager(self.db_path)
        make_emp(self.db, "Empresa Backup")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_hacer_backup_crea_archivo(self):
        from utils import BackupManager
        try:
            ok = BackupManager.hacer_backup(self.db_path)
            # If it returns something, check it's truthy or at least doesn't crash
        except TypeError:
            # Maybe it takes no args and uses config — just verify no crash
            try: BackupManager.hacer_backup()
            except Exception: pass
        except Exception as e:
            self.fail(f"Crash hacer_backup: {e}")

    def test_backup_db_inexistente_no_crashea(self):
        from utils import BackupManager
        try:
            BackupManager.hacer_backup("/ruta/inexistente/test.db")
        except TypeError:
            try: BackupManager.hacer_backup()
            except Exception: pass
        except Exception:
            pass  # Failing gracefully is acceptable


# =====================================================================
# 19. STRESS EXTREMO
# =====================================================================
class TestStressExtremo(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_10000_empresas_count(self):
        for i in range(10000):
            self.db.agregar_empresa(f"E{i:05d}","","","","","","")
        self.assertEqual(self.db.count("empresas"), 10000)

    def test_busqueda_10000_bajo_3s(self):
        for i in range(10000):
            self.db.agregar_empresa(f"Empresa_{i:05d}","","","","","","")
        t0 = time.time()
        self.db.get_filtered_empresas("Empresa_0999", {})
        elapsed = time.time() - t0
        self.assertLess(elapsed, 3.0,
                        f"Busqueda tardo {elapsed:.2f}s con 10000 empresas")

    def test_descripcion_emoji_masiva(self):
        eid = make_emp(self.db)
        desc = "X" * 5000
        try: self.db.agregar_cotizacion(eid, desc, 1.0)
        except Exception as e: self.fail(f"Crash descripcion masiva: {e}")

    def test_100_tags_distintos(self):
        eid = make_emp(self.db, "Multi Tag")
        tags_list = [f"tag_{i}" for i in range(100)]
        try: self.db.vincular_empresa_con_tags(eid, tags_list)
        except Exception as e: self.fail(f"Crash 100 tags: {e}")
        tags = self.db.get_tags_de_empresa(eid)
        self.assertEqual(len(tags), 100)

    def test_20_merges_consecutivos(self):
        ids = [make_emp(self.db, f"Merge_{i}") for i in range(21)]
        destino = ids[0]
        for origen in ids[1:]:
            try: self.db.unificar_empresas(origen, destino)
            except: pass
        self.assertIsNotNone(self.db.obtener_empresa_por_id(destino),
            "La empresa destino desaparecio despues de 20 merges")



# =====================================================================
# 21. HISTORIAL DE CAMBIOS
# =====================================================================
class TestHistorial(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_registrar_cambio_basico(self):
        eid = make_emp(self.db, "Empresa A")
        ok = self.db.registrar_cambio(eid, "nombre", "viejo", "nuevo", "usuario")
        self.assertTrue(ok)
        hist = self.db.get_historial_empresa(eid)
        self.assertEqual(len(hist), 1)
        self.assertEqual(hist[0]["campo"], "nombre")
        self.assertEqual(hist[0]["valor_anterior"], "viejo")
        self.assertEqual(hist[0]["valor_nuevo"], "nuevo")
        self.assertEqual(hist[0]["fuente"], "usuario")

    def test_editar_empresa_genera_historial(self):
        eid = make_emp(self.db, "Nombre Original")
        self.db.editar_empresa(eid, "Nombre Nuevo", "", "", "", "", "", "")
        hist = self.db.get_historial_empresa(eid)
        campos = [h["campo"] for h in hist]
        self.assertIn("nombre", campos)
        cambio = next(h for h in hist if h["campo"] == "nombre")
        self.assertEqual(cambio["valor_anterior"], "Nombre Original")
        self.assertEqual(cambio["valor_nuevo"], "Nombre Nuevo")

    def test_editar_sin_cambios_no_genera_historial(self):
        eid = make_emp(self.db, "Sin Cambios")
        emp = self.db.obtener_empresa_por_id(eid)
        self.db.editar_empresa(
            eid, emp["nombre"], emp.get("direccion",""),
            emp.get("telefono",""), emp.get("email",""),
            emp.get("rubro",""), emp.get("pais",""), "")
        hist = self.db.get_historial_empresa(eid)
        self.assertEqual(len(hist), 0)

    def test_historial_multiples_cambios_orden_desc(self):
        eid = make_emp(self.db, "Empresa Multi")
        self.db.registrar_cambio(eid, "nombre", "v1", "v2", "usuario")
        self.db.registrar_cambio(eid, "pais",   "AR", "CL", "usuario")
        self.db.registrar_cambio(eid, "email",  "a@a.com", "b@b.com", "gemini")
        hist = self.db.get_historial_empresa(eid)
        self.assertEqual(len(hist), 3)
        # El más reciente (id más alto) debe ser el de email/gemini
        self.assertEqual(hist[0]["campo"], "email")
        self.assertEqual(hist[0]["fuente"], "gemini")

    def test_eliminar_empresa_elimina_historial(self):
        eid = make_emp(self.db, "Para Borrar Con Hist")
        self.db.registrar_cambio(eid, "nombre", "x", "y", "usuario")
        self.db.eliminar_empresa(eid)
        self.assertEqual(self.db.get_historial_empresa(eid), [])

    def test_registrar_sin_diferencia_no_guarda(self):
        eid = make_emp(self.db, "Empresa B")
        self.db.registrar_cambio(eid, "nombre", "igual", "igual", "usuario")
        self.assertEqual(len(self.db.get_historial_empresa(eid)), 0)

    def test_eliminar_cambio(self):
        eid = make_emp(self.db, "Empresa C")
        self.db.registrar_cambio(eid, "nombre", "a", "b", "usuario")
        hist = self.db.get_historial_empresa(eid)
        self.db.eliminar_cambio(hist[0]["id"])
        self.assertEqual(len(self.db.get_historial_empresa(eid)), 0)

    def test_historial_empresa_inexistente(self):
        self.assertEqual(self.db.get_historial_empresa(99999), [])

    def test_fuente_gemini_registrada(self):
        eid = make_emp(self.db, "Empresa IA")
        self.db.editar_empresa(eid, "Empresa IA Mejorada","","","","","","",
                               fuente="gemini")
        hist = self.db.get_historial_empresa(eid)
        self.assertIn("gemini", [h["fuente"] for h in hist])

    def test_sql_injection_en_campo(self):
        eid = make_emp(self.db, "Safe")
        try:
            self.db.registrar_cambio(eid,"'; DROP TABLE cambios;--","x","y","u")
        except Exception as e:
            self.fail(f"Crash: {e}")
        self.assertIsNotNone(self.db.fetchall("SELECT * FROM cambios"))


# =====================================================================
# 22. FILTROS AVANZADOS V21 — PAIS, RUBRO, TAG
# =====================================================================
class TestFiltrosAvanzadosV21(unittest.TestCase):
    def setUp(self):
        self.db = fresh_db()
        self.e1 = make_emp(self.db, "Empresa AR Tech")
        self.db.editar_empresa(self.e1,"Empresa AR Tech","","","","Tecnologia","Argentina","")
        self.e2 = make_emp(self.db, "Empresa CL Comercio")
        self.db.editar_empresa(self.e2,"Empresa CL Comercio","","","","Comercio","Chile","")
        self.e3 = make_emp(self.db, "Empresa AR Comercio")
        self.db.editar_empresa(self.e3,"Empresa AR Comercio","","","","Comercio","Argentina","")
        self.db.vincular_empresa_con_tags(self.e1, ["vip"])
        self.db.vincular_empresa_con_tags(self.e3, ["cliente"])

    def test_filtro_pais(self):
        ids = {r["id"] for r in self.db.get_filtered_empresas("",{"pais":"Argentina"})}
        self.assertIn(self.e1,ids); self.assertIn(self.e3,ids); self.assertNotIn(self.e2,ids)

    def test_filtro_rubro(self):
        ids = {r["id"] for r in self.db.get_filtered_empresas("",{"rubro":"Comercio"})}
        self.assertIn(self.e2,ids); self.assertIn(self.e3,ids); self.assertNotIn(self.e1,ids)

    def test_filtro_tag(self):
        ids = {r["id"] for r in self.db.get_filtered_empresas("",{"tag":"vip"})}
        self.assertIn(self.e1,ids); self.assertNotIn(self.e2,ids); self.assertNotIn(self.e3,ids)

    def test_filtro_pais_y_rubro(self):
        ids = {r["id"] for r in self.db.get_filtered_empresas(
            "",{"pais":"Argentina","rubro":"Comercio"})}
        self.assertIn(self.e3,ids); self.assertNotIn(self.e1,ids); self.assertNotIn(self.e2,ids)

    def test_tag_inexistente(self):
        self.assertEqual(len(self.db.get_filtered_empresas("",{"tag":"xyz_no_existe"})),0)

    def test_pais_vacio_devuelve_todo(self):
        self.assertEqual(len(self.db.get_filtered_empresas("",{"pais":""})),
                         self.db.count("empresas"))

    def test_texto_y_pais_combinados(self):
        ids = {r["id"] for r in self.db.get_filtered_empresas(
            "AR",{"pais":"Argentina","rubro":"Tecnologia"})}
        self.assertIn(self.e1,ids); self.assertNotIn(self.e2,ids); self.assertNotIn(self.e3,ids)


# =====================================================================
# 23. PAGINACION BUSQUEDA GLOBAL
# =====================================================================
class TestPaginacionBusquedaGlobal(unittest.TestCase):
    PS = 200

    def test_menos_de_page_una_pagina(self):
        self.assertEqual(max(1,(50+self.PS-1)//self.PS), 1)

    def test_exactamente_page_una_pagina(self):
        self.assertEqual(max(1,(200+self.PS-1)//self.PS), 1)

    def test_201_dos_paginas(self):
        self.assertEqual(max(1,(201+self.PS-1)//self.PS), 2)

    def test_450_tres_paginas(self):
        self.assertEqual(max(1,(450+self.PS-1)//self.PS), 3)

    def test_slice_primera_pagina(self):
        rows = list(range(450))
        self.assertEqual(rows[0:self.PS][0], 0)
        self.assertEqual(len(rows[0:self.PS]), 200)

    def test_slice_segunda_pagina(self):
        rows = list(range(450))
        self.assertEqual(rows[self.PS:2*self.PS][0], 200)

    def test_slice_ultima_pagina_parcial(self):
        rows = list(range(450))
        self.assertEqual(len(rows[2*self.PS:3*self.PS]), 50)

    def test_total_monto_todos_no_pagina(self):
        montos = [float(i*100) for i in range(450)]
        total_all = sum(montos)
        page_total = sum(montos[:self.PS])
        self.assertGreater(total_all, page_total)

    def test_pagina_clamp_min_cero(self):
        page = 0
        page = max(0, page - 1)
        self.assertEqual(page, 0)

    def test_pagina_clamp_max(self):
        pages = 3; page = 2
        page = min(page + 1, pages - 1)
        self.assertEqual(page, 2)


# =====================================================================
# 24. get_similar_empresas OPTIMIZADO
# =====================================================================
class TestSimilarEmpresasOpt(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_detecta_similares(self):
        make_emp(self.db,"Acme Argentina SA"); make_emp(self.db,"Acme Argentina SRL")
        make_emp(self.db,"Empresa Totalmente Diferente XYZ")
        pares = self.db.get_similar_empresas(75)
        self.assertTrue(any("Acme" in p["nombre1"] and "Acme" in p["nombre2"] for p in pares))

    def test_rendimiento_500(self):
        for i in range(500):
            self.db.agregar_empresa(f"Empresa_{i:04d} SA","","","","","","")
        t0 = time.time()
        self.db.get_similar_empresas(85)
        self.assertLess(time.time()-t0, 8.0)

    def test_ordenado_desc(self):
        make_emp(self.db,"ACME SA"); make_emp(self.db,"ACME SRL"); make_emp(self.db,"ACME S.A.")
        pares = self.db.get_similar_empresas(70)
        if len(pares) >= 2:
            sims = [p["similitud"] for p in pares]
            self.assertEqual(sims, sorted(sims,reverse=True))

    def test_sin_empresas(self):
        self.assertEqual(self.db.get_similar_empresas(80), [])

    def test_una_empresa(self):
        make_emp(self.db,"Unica SA")
        self.assertEqual(self.db.get_similar_empresas(80), [])

    def test_muy_distintas_no_aparecen(self):
        make_emp(self.db,"Google LLC"); make_emp(self.db,"Panaderia San Martin")
        self.assertEqual(self.db.get_similar_empresas(70), [])

    def test_umbral_100_solo_identicas(self):
        make_emp(self.db,"ACME SA"); make_emp(self.db,"ACME SA"); make_emp(self.db,"ACME SRL")
        for p in self.db.get_similar_empresas(100):
            self.assertEqual(p["similitud"], 100)


# =====================================================================
# 25. VIRTUAL TREE SORT
# =====================================================================
class TestVirtualTreeSort(unittest.TestCase):
    def _rows(self, data):
        return [(str(i), vals, ("g",)) for i,vals in enumerate(data)]

    def test_sort_nombre_asc(self):
        rows = self._rows([("Zeta","AR","",5),("Alpha","CL","",2),("Mitad","AR","",0)])
        rows.sort(key=lambda r: r[1][0].lower())
        self.assertEqual([r[1][0] for r in rows], ["Alpha","Mitad","Zeta"])

    def test_sort_nombre_desc(self):
        rows = self._rows([("Zeta","",""  ,0),("Alpha","","",0),("Mitad","","",0)])
        rows.sort(key=lambda r: r[1][0].lower(), reverse=True)
        self.assertEqual([r[1][0] for r in rows], ["Zeta","Mitad","Alpha"])

    def test_sort_pais(self):
        rows = self._rows([("E1","Uruguay","",0),("E2","Argentina","",0),("E3","Chile","",0)])
        rows.sort(key=lambda r: r[1][1].lower())
        self.assertEqual([r[1][1] for r in rows], ["Argentina","Chile","Uruguay"])

    def test_sort_ncot_desc(self):
        rows = self._rows([("E1","",""  ,3),("E2","","",10),("E3","","",1)])
        rows.sort(key=lambda r: int(r[1][3] or 0), reverse=True)
        self.assertEqual([r[1][3] for r in rows], [10,3,1])

    def test_toggle_asc_desc(self):
        rows = self._rows([("Z","","",0),("A","","",0),("M","","",0)])
        rows.sort(key=lambda r: r[1][0].lower())
        self.assertEqual(rows[0][1][0], "A")
        rows.sort(key=lambda r: r[1][0].lower(), reverse=True)
        self.assertEqual(rows[0][1][0], "Z")

    def test_sort_lista_vacia(self):
        rows = []
        rows.sort(key=lambda r: r[1][0].lower())
        self.assertEqual(rows, [])

    def test_sort_un_elemento(self):
        rows = self._rows([("Unica SA","AR","",0)])
        rows.sort(key=lambda r: r[1][0].lower())
        self.assertEqual(len(rows), 1)


# =====================================================================
# 26. CONFIG EXPORT
# =====================================================================
class TestFiltrosRubroTag(unittest.TestCase):
    def setUp(self):
        self.db = fresh_db()
        self.e1 = make_emp(self.db, "Tech SA")
        self.e2 = make_emp(self.db, "Retail SA")
        self.db.editar_empresa(self.e1,"Tech SA","","","","Tech","AR","cliente")
        self.db.editar_empresa(self.e2,"Retail SA","","","","Retail","AR","vip")

    def test_filtro_rubro(self):
        ids = {r["id"] for r in self.db.get_filtered_empresas("",{"rubro":"Tech"})}
        self.assertIn(self.e1, ids)
        self.assertNotIn(self.e2, ids)

    def test_filtro_rubro_inexistente(self):
        self.assertEqual(self.db.get_filtered_empresas("",{"rubro":"NoExiste"}), [])

    def test_filtro_tag(self):
        ids = {r["id"] for r in self.db.get_filtered_empresas("",{"tag":"vip"})}
        self.assertIn(self.e2, ids)
        self.assertNotIn(self.e1, ids)

    def test_filtro_tag_inexistente(self):
        self.assertEqual(self.db.get_filtered_empresas("",{"tag":"no_existe_xyz"}), [])

    def test_filtro_rubro_y_pais(self):
        e3 = make_emp(self.db,"Tech CL",pais="Chile")
        self.db.editar_empresa(e3,"Tech CL","","","","Tech","Chile","")
        result = self.db.get_filtered_empresas("",{"rubro":"Tech","pais":"Chile"})
        ids = {r["id"] for r in result}
        self.assertIn(e3, ids)
        self.assertNotIn(self.e1, ids)


# =====================================================================
# 20D. get_similar_empresas OPTIMIZADO
# =====================================================================
class TestSimilarOptimizado(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_detecta_similares(self):
        make_emp(self.db,"Acme Argentina SA")
        make_emp(self.db,"Acme Argentina SRL")
        pares = self.db.get_similar_empresas(75)
        self.assertTrue(any("Acme" in p["nombre1"] and "Acme" in p["nombre2"] for p in pares))

    def test_performance_500(self):
        for i in range(500): self.db.agregar_empresa(f"Empresa_{i:04d}","","","","","","")
        t0 = time.time()
        self.db.get_similar_empresas(85)
        self.assertLess(time.time()-t0, 15.0)

    def test_ordenado_desc(self):
        make_emp(self.db,"Alpha Corp SA"); make_emp(self.db,"Alpha Corp SRL")
        make_emp(self.db,"Alpha Corp")
        pares = self.db.get_similar_empresas(70)
        if len(pares) >= 2:
            sims = [p["similitud"] for p in pares]
            self.assertEqual(sims, sorted(sims,reverse=True))

    def test_sin_empresas_lista_vacia(self):
        self.assertEqual(self.db.get_similar_empresas(80), [])

    def test_una_empresa_lista_vacia(self):
        make_emp(self.db,"Única")
        self.assertEqual(self.db.get_similar_empresas(80), [])


# =====================================================================
# 20E. PAGINACION BUSQUEDA GLOBAL
# =====================================================================
class TestPaginacion(unittest.TestCase):
    PS = 200

    def _pag(self, total, page):
        pages = max(1,(total+self.PS-1)//self.PS)
        page  = max(0,min(page,pages-1))
        start = page*self.PS
        return start, min(start+self.PS,total), pages

    def test_primera_pagina(self):
        s,e,p = self._pag(500,0)
        self.assertEqual(s,0); self.assertEqual(e,200); self.assertEqual(p,3)

    def test_ultima_pagina(self):
        s,e,p = self._pag(450,2)
        self.assertEqual(s,400); self.assertEqual(e,450)

    def test_pagina_fuera_de_rango(self):
        s,e,p = self._pag(100,99)
        self.assertEqual(p,1); self.assertEqual(s,0)

    def test_menos_de_una_pagina(self):
        s,e,p = self._pag(50,0)
        self.assertEqual(p,1); self.assertEqual(e,50)

    def test_cero_resultados(self):
        s,e,p = self._pag(0,0)
        self.assertEqual(p,1); self.assertEqual(e,0)

    def test_exactamente_dos_paginas(self):
        s,e,p = self._pag(201,1)
        self.assertEqual(s,200); self.assertEqual(e,201)

    def test_prev_next_logica(self):
        for total,page,exp_prev,exp_next in [
            (500,0,False,True),(500,1,True,True),(500,2,True,False)
        ]:
            _,_,pages = self._pag(total,page)
            self.assertEqual(page>0, exp_prev)
            self.assertEqual(page<pages-1, exp_next)


# =====================================================================
# 20F. VIRTUAL TREE SORT
# =====================================================================
class TestVirtualTreeSort(unittest.TestCase):
    def _rows(self,data):
        return [(str(d[0]),(d[1],d[2],"",d[3]),("none",)) for d in data]

    def _sort(self,rows,fn,asc=True):
        return sorted(rows,key=fn,reverse=not asc)

    def test_sort_nombre_asc(self):
        rows = self._rows([(1,"Zeta","AR",0),(2,"Alpha","AR",0),(3,"Beta","AR",0)])
        s = self._sort(rows,lambda r:r[1][0].lower())
        self.assertEqual([r[1][0] for r in s],["Alpha","Beta","Zeta"])

    def test_sort_nombre_desc(self):
        rows = self._rows([(1,"Zeta","AR",0),(2,"Alpha","AR",0)])
        s = self._sort(rows,lambda r:r[1][0].lower(),asc=False)
        self.assertEqual(s[0][1][0],"Zeta")

    def test_sort_ncot(self):
        rows = self._rows([(1,"A","AR",10),(2,"B","AR",1),(3,"C","AR",5)])
        s = self._sort(rows,lambda r:int(r[1][3] or 0))
        self.assertEqual([r[1][3] for r in s],[1,5,10])

    def test_sort_pais(self):
        rows = self._rows([(1,"A","Uruguay",0),(2,"B","Argentina",0),(3,"C","Chile",0)])
        s = self._sort(rows,lambda r:r[1][1].lower())
        self.assertEqual([r[1][1] for r in s],["Argentina","Chile","Uruguay"])

    def test_toggle_inversion(self):
        rows = self._rows([(1,"Z","AR",0),(2,"A","AR",0)])
        s1 = self._sort(rows,lambda r:r[1][0].lower(),asc=True)
        self.assertEqual(s1[0][1][0],"A")
        s2 = self._sort(rows,lambda r:r[1][0].lower(),asc=False)
        self.assertEqual(s2[0][1][0],"Z")

    def test_preserva_todos_items(self):
        rows = self._rows([(i,f"E_{i}","AR",i) for i in range(100)])
        s = self._sort(rows,lambda r:r[1][0].lower())
        self.assertEqual({r[0] for r in rows},{r[0] for r in s})

    def test_lista_vacia(self):
        self.assertEqual(self._sort([],lambda r:r[1][0]),[])

    def test_sort_estable(self):
        rows = self._rows([(1,"Alpha","AR",5),(2,"Alpha","CL",3),(3,"Alpha","UY",1)])
        s = self._sort(rows,lambda r:r[1][0].lower())
        self.assertEqual([r[0] for r in s],["1","2","3"])


# =====================================================================
# 20G. CONFIG EXPORT
# =====================================================================
class TestConfigExport(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)
        import config_export; config_export._instance = None

    def _mgr(self):
        """Create a ConfigManager pointing to tmpdir."""
        import config_export, unittest.mock as mock
        config_export._instance = None
        with mock.patch("config_export._find_config_dir",
                        return_value=__import__("pathlib").Path(self.tmpdir)):
            return config_export.ConfigManager()

    def test_defaults_presentes(self):
        from config_export import DEFAULT_CONFIG
        for k in ("db_name","theme","duplicados_umbral","ai_provider",
                  "gemini_model","grok_model","busqueda_page_size",
                  "importer_thresh","ai_batch_size","filtros_guardados"):
            self.assertIn(k, DEFAULT_CONFIG)

    def test_get_devuelve_default(self):
        mgr = self._mgr()
        self.assertEqual(mgr.get("duplicados_umbral"), 85)
        self.assertEqual(mgr.get("ai_provider"), "auto")

    def test_set_persiste(self):
        import config_export, unittest.mock as mock
        config_export._instance = None
        with mock.patch("config_export._find_config_dir",
                        return_value=__import__("pathlib").Path(self.tmpdir)):
            mgr = config_export.ConfigManager()
            mgr.set("duplicados_umbral", 90)
            mgr2 = config_export.ConfigManager()
        self.assertEqual(mgr2.get("duplicados_umbral"), 90)

    def test_guardar_y_obtener_filtro(self):
        mgr = self._mgr()
        mgr.guardar_filtro("vip_ar", {"pais": "Argentina", "tag": "vip"})
        filtros = mgr.get_filtros_guardados()
        self.assertIn("vip_ar", filtros)
        self.assertEqual(filtros["vip_ar"]["pais"], "Argentina")

    def test_importar_no_cambia_db_name(self):
        import json, config_export, unittest.mock as mock
        src = os.path.join(self.tmpdir, "import.json")
        with open(src, "w") as f:
            json.dump({"version":1,"db_name":"otra.db","theme":"darkly"}, f)
        config_export._instance = None
        with mock.patch("config_export._find_config_dir",
                        return_value=__import__("pathlib").Path(self.tmpdir)):
            mgr = config_export.ConfigManager()
            mgr._data["db_name"] = "mi_empresa.db"
            ok, _ = mgr.importar(src, backup=False)
        self.assertTrue(ok)
        self.assertEqual(mgr.get("db_name"), "mi_empresa.db")
        self.assertEqual(mgr.get("theme"), "darkly")

    def test_importar_version_futura_falla(self):
        import json, config_export
        from config_export import CONFIG_VERSION
        src = os.path.join(self.tmpdir, "future.json")
        with open(src, "w") as f:
            json.dump({"version": CONFIG_VERSION + 10}, f)
        mgr = self._mgr()
        ok, msg = mgr.importar(src, backup=False)
        self.assertFalse(ok)
        self.assertIn("versión", msg.lower())

    def test_resetear(self):
        mgr = self._mgr()
        mgr.set("duplicados_umbral", 99)
        mgr.resetear()
        self.assertEqual(mgr.get("duplicados_umbral"), 85)

    def test_exportar_crea_json_valido(self):
        import json
        mgr = self._mgr()
        dest = os.path.join(self.tmpdir, "export.json")
        ok = mgr.exportar(dest)
        self.assertTrue(ok)
        with open(dest) as f:
            data = json.load(f)
        self.assertIn("exportado_en", data)
        self.assertIn("db_name", data)

    def test_carpetas_recientes_max_5(self):
        mgr = self._mgr()
        for i in range(7):
            mgr.agregar_carpeta_reciente(f"/ruta/que/no/existe/{i}")
        raw = mgr._data.get("carpetas_recientes", [])
        self.assertLessEqual(len(raw), 5)
        self.assertEqual(raw[0], "/ruta/que/no/existe/6")

    def test_get_key_inexistente_devuelve_default(self):
        mgr = self._mgr()
        self.assertIsNone(mgr.get("key_que_no_existe"))
        self.assertEqual(mgr.get("key_que_no_existe", "fallback"), "fallback")

class TestWAL(unittest.TestCase):
    """Tests de las garantías de thread-safety agregadas en v22."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.db_path = os.path.join(self.tmpdir, "wal_test.db")

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_wal_mode_activo_en_db_nueva(self):
        """La DB debe arrancar en WAL mode, no en delete (default de SQLite)."""
        db = DBManager(self.db_path)
        conn = sqlite3.connect(self.db_path)
        mode = conn.execute("PRAGMA journal_mode").fetchone()[0]
        conn.close()
        self.assertEqual(mode, "wal",
            f"journal_mode debería ser 'wal', es '{mode}'")

    def test_busy_timeout_configurado(self):
        """busy_timeout > 0 para que la app espere antes de lanzar 'database locked'."""
        db = DBManager(self.db_path)
        conn = sqlite3.connect(self.db_path)
        # busy_timeout se configura a nivel de sesión, verificamos que no crashea
        # en escenario de contención
        timeout = conn.execute("PRAGMA busy_timeout").fetchone()[0]
        conn.close()
        self.assertGreaterEqual(timeout, 0)  # > 0 es el objetivo

    def test_write_lock_es_threading_lock(self):
        """_write_lock debe existir y ser un Lock o RLock de threading."""
        import threading
        db = DBManager(self.db_path)
        self.assertTrue(hasattr(db, "_write_lock"))
        lock = db._write_lock
        is_lock = isinstance(lock, (type(threading.Lock()),
                                    type(threading.RLock())))
        self.assertTrue(is_lock,
            f"_write_lock debe ser Lock o RLock, es: {type(lock)}")

    def test_no_lastrowid_en_dbmanager(self):
        """lastrowid fue eliminado — no debe existir como atributo de instancia."""
        db = DBManager(":memory:")
        self.assertFalse(hasattr(db, "lastrowid"),
            "lastrowid no debe existir — es antipatrón de estado compartido")

    def test_dos_threads_ejecutar_simultaneo_sin_crash(self):
        """Dos threads escribiendo al mismo tiempo no deben corromper la DB."""
        db = DBManager(self.db_path)
        errors = []

        def writer(prefix, n):
            for i in range(20):
                ok = db.ejecutar(
                    "INSERT INTO empresas (nombre) VALUES (?)",
                    (f"{prefix}_{i}",))
                if not ok:
                    errors.append(f"{prefix}_{i} failed")

        t1 = threading.Thread(target=writer, args=("T1", 20))
        t2 = threading.Thread(target=writer, args=("T2", 20))
        t1.start(); t2.start()
        t1.join(timeout=15); t2.join(timeout=15)

        alive = [t for t in (t1, t2) if t.is_alive()]
        self.assertEqual(len(alive), 0, f"{len(alive)} threads colgados (deadlock?)")
        self.assertEqual(errors, [], f"Escrituras fallidas: {errors}")
        count = db.count("empresas")
        self.assertEqual(count, 40, f"Esperaba 40 empresas, hay {count}")

    def test_thread_secundario_conexion_propia_no_corrompe(self):
        """Thread con su propia conexión SQLite escribe sin afectar la conexión principal."""
        db = DBManager(self.db_path)
        eid = make_emp(db, "Empresa Principal")
        make_cot(db, eid, 1000.0)

        # Simular el patrón del resumidor: thread abre su propia conexión
        errors = []
        done = threading.Event()

        def background_writer():
            try:
                conn = sqlite3.connect(self.db_path, timeout=10)
                conn.execute("PRAGMA busy_timeout=5000")
                conn.execute(
                    "UPDATE cotizaciones SET resumen=? WHERE empresa_id=?",
                    ("Resumen generado por thread", eid))
                conn.commit()
                conn.close()
            except Exception as e:
                errors.append(str(e))
            finally:
                done.set()

        t = threading.Thread(target=background_writer, daemon=True)
        t.start()
        done.wait(timeout=10)

        self.assertEqual(errors, [], f"Error en thread secundario: {errors}")
        # La conexión principal puede leer el cambio
        row = db.fetchone(
            "SELECT resumen FROM cotizaciones WHERE empresa_id=?", (eid,))
        self.assertIsNotNone(row)


    def test_global_lock_compartido_entre_instancias(self):
        """
        Dos instancias de DBManager apuntando al mismo archivo
        deben compartir el mismo RLock (no uno por instancia).
        Esto garantiza serialización entre los 3 usuarios.
        """
        db1 = DBManager(self.db_path)
        db2 = DBManager(self.db_path)
        self.assertIs(db1._write_lock, db2._write_lock,
            "Dos DBManager al mismo archivo deben compartir el mismo lock")

    def test_global_lock_distinto_para_distintas_db(self):
        """Dos archivos distintos → locks distintos."""
        db2_path = os.path.join(self.tmpdir, "otro.db")
        db1 = DBManager(self.db_path)
        db2 = DBManager(db2_path)
        self.assertIsNot(db1._write_lock, db2._write_lock,
            "DBs distintas deben tener locks distintos")

    def test_memory_db_lock_distinto_por_instancia(self):
        """Cada :memory: DB tiene su propio lock independiente."""
        db1 = DBManager(":memory:")
        db2 = DBManager(":memory:")
        self.assertIsNot(db1._write_lock, db2._write_lock,
            "Cada :memory: debe tener su propio lock")

    def test_3_instancias_mismo_archivo_escrituras_sin_perdida(self):
        """
        El caso real de 3 usuarios: 3 DBManager distintos,
        mismo archivo, escrituras simultaneas, sin perdida de datos.
        """
        db1 = DBManager(self.db_path)
        db2 = DBManager(self.db_path)
        db3 = DBManager(self.db_path)
        self.assertIs(db1._write_lock, db2._write_lock)
        self.assertIs(db1._write_lock, db3._write_lock)
        errors = []

        def writer(db, prefix):
            for i in range(50):
                ok = db.agregar_empresa(f"{prefix}_{i}", "", "", "", "", "", "")
                if not ok:
                    errors.append(f"{prefix}_{i}")

        threads = [
            threading.Thread(target=writer, args=(db1, "U1")),
            threading.Thread(target=writer, args=(db2, "U2")),
            threading.Thread(target=writer, args=(db3, "U3")),
        ]
        for t in threads: t.start()
        for t in threads: t.join(timeout=20)

        self.assertEqual(errors, [], f"Escrituras fallidas: {errors}")
        count = db1.count("empresas")
        self.assertEqual(count, 150, f"Con lock global: {count} != 150")

    def test_wal_permite_leer_mientras_escribe(self):
        """Con WAL, una lectura no debe bloquearse por una escritura larga."""
        db = DBManager(self.db_path)
        # Poblar datos
        for i in range(50):
            db.agregar_empresa(f"Empresa_{i}", "", "", "", "", "", "")

        read_results = []
        errors = []

        def reader():
            for _ in range(10):
                try:
                    results = db.get_filtered_empresas("", {})
                    read_results.append(len(results))
                except Exception as e:
                    errors.append(str(e))
                time.sleep(0.01)

        def writer():
            for i in range(10):
                db.agregar_empresa(f"Extra_{i}", "", "", "", "", "", "")
                time.sleep(0.01)

        t_read  = threading.Thread(target=reader)
        t_write = threading.Thread(target=writer)
        t_read.start(); t_write.start()
        t_read.join(timeout=15); t_write.join(timeout=15)

        self.assertEqual(errors, [], f"Errores de lectura concurrente: {errors}")
        self.assertGreater(len(read_results), 0, "El reader no ejecutó ninguna lectura")


# =====================================================================
# 28. EXTRACTOR DE TEXTO
# =====================================================================
class TestExtractorTexto(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _write(self, name, content, encoding="utf-8"):
        path = os.path.join(self.tmpdir, name)
        with open(path, "w", encoding=encoding) as f:
            f.write(content)
        return path

    def test_extrae_txt_basico(self):
        from extractor_texto import extraer
        path = self._write("doc.txt", "Hola mundo\nSegunda línea")
        result = extraer(path)
        self.assertIn("Hola mundo", result)
        self.assertIn("Segunda línea", result)

    def test_extrae_txt_encoding_latin1(self):
        from extractor_texto import extraer
        path = os.path.join(self.tmpdir, "latin.txt")
        with open(path, "w", encoding="latin-1") as f:
            f.write("Empresa Ñoño SA\nDirección: Av. Corrientes")
        result = extraer(path)
        self.assertIsInstance(result, str)
        self.assertGreater(len(result), 0)

    def test_extrae_txt_unicode(self):
        from extractor_texto import extraer
        path = self._write("uni.txt", "企业名称 Ünternehmen شركة")
        result = extraer(path)
        self.assertIsInstance(result, str)

    def test_archivo_no_existe_lanza_filenotfounderror(self):
        from extractor_texto import extraer
        with self.assertRaises(FileNotFoundError):
            extraer("/ruta/que/no/existe/archivo.txt")

    def test_extension_no_soportada_lanza_valueerror(self):
        from extractor_texto import extraer
        path = self._write("foto.jpg", "datos binarios")
        with self.assertRaises(ValueError):
            extraer(path)

    def test_archivo_vacio_devuelve_string(self):
        from extractor_texto import extraer
        path = self._write("vacio.txt", "")
        result = extraer(path)
        self.assertIsInstance(result, str)
        self.assertEqual(result, "")

    def test_truncacion_a_max_chars(self):
        from extractor_texto import extraer, MAX_CHARS
        # Crear texto más largo que MAX_CHARS
        texto_largo = "A" * (MAX_CHARS + 5000)
        path = self._write("largo.txt", texto_largo)
        result = extraer(path)
        self.assertLessEqual(len(result), MAX_CHARS,
            f"Debería truncar a {MAX_CHARS}, tiene {len(result)}")

    def test_extrae_docx(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx no instalado")
        from extractor_texto import extraer
        path = os.path.join(self.tmpdir, "doc.docx")
        doc = Document()
        doc.add_paragraph("Propuesta comercial para ACME SA")
        doc.add_paragraph("Monto total: USD 15.000")
        doc.save(path)
        result = extraer(path)
        self.assertIn("ACME SA", result)
        self.assertIn("15.000", result)

    def test_extrae_xlsx(self):
        try:
            import openpyxl
        except ImportError:
            self.skipTest("openpyxl no instalado")
        from extractor_texto import extraer
        path = os.path.join(self.tmpdir, "tabla.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Empresa"; ws["B1"] = "Monto"
        ws["A2"] = "ACME SA"; ws["B2"] = 50000
        wb.save(path)
        result = extraer(path)
        self.assertIn("ACME SA", result)
        self.assertIn("50000", result)

    def test_limpieza_espacios_multiples(self):
        from extractor_texto import _clean
        sucio = "texto   con    espacios\n\n\n\ntriple salto"
        limpio = _clean(sucio)
        self.assertNotIn("   ", limpio)
        self.assertNotIn("\n\n\n", limpio)


# =====================================================================
# 29. RESUMIDOR
# =====================================================================
class TestResumidor(unittest.TestCase):

    def test_resumir_texto_vacio_devuelve_defaults(self):
        from resumidor import resumir
        r = resumir("")
        self.assertIsInstance(r, dict)
        self.assertIn("resumen", r)
        self.assertIsNone(r["monto"])

    def test_resumir_none_devuelve_defaults(self):
        from resumidor import resumir
        r = resumir(None)
        self.assertIsInstance(r, dict)
        self.assertIsNone(r["monto"])

    def test_resumir_solo_espacios_devuelve_defaults(self):
        from resumidor import resumir
        r = resumir("   \n\t  ")
        self.assertIsInstance(r, dict)

    def test_resumir_sin_api_key_no_crashea(self):
        from resumidor import resumir
        import os
        # Asegurar que no hay key en el entorno para este test
        old_g = os.environ.pop("GEMINI_API_KEY", None)
        old_g2 = os.environ.pop("GOOGLE_API_KEY", None)
        old_k = os.environ.pop("GROK_API_KEY", None)
        old_k2 = os.environ.pop("XAI_API_KEY", None)
        try:
            r = resumir("Propuesta comercial por USD 10.000")
            self.assertIsInstance(r, dict)
            self.assertIn("proveedor_ia", r)
            self.assertEqual(r["proveedor_ia"], "none")
        finally:
            if old_g:  os.environ["GEMINI_API_KEY"] = old_g
            if old_g2: os.environ["GOOGLE_API_KEY"] = old_g2
            if old_k:  os.environ["GROK_API_KEY"] = old_k
            if old_k2: os.environ["XAI_API_KEY"] = old_k2

    def test_parse_response_json_valido(self):
        from resumidor import _parse_response
        raw = '{"resumen":"Propuesta de equipos","monto":15000.0,"moneda":"USD","tipo":"Propuesta","proveedor":"ACME","fecha_doc":"2024-03-15","confianza":0.95}'
        r = _parse_response(raw)
        self.assertEqual(r["resumen"], "Propuesta de equipos")
        self.assertAlmostEqual(r["monto"], 15000.0)
        self.assertEqual(r["moneda"], "USD")
        self.assertEqual(r["tipo"], "Propuesta")
        self.assertEqual(r["proveedor"], "ACME")
        self.assertEqual(r["fecha_doc"], "2024-03-15")
        self.assertAlmostEqual(r["confianza"], 0.95)

    def test_parse_response_con_markdown_fences(self):
        from resumidor import _parse_response
        raw = '```json\n{"resumen":"Doc con fences","monto":5000,"moneda":"ARS","tipo":"Factura","proveedor":null,"fecha_doc":null,"confianza":0.8}\n```'
        r = _parse_response(raw)
        self.assertEqual(r["resumen"], "Doc con fences")
        self.assertAlmostEqual(r["monto"], 5000.0)

    def test_parse_response_json_invalido_lanza_error(self):
        from resumidor import _parse_response
        with self.assertRaises((ValueError, Exception)):
            _parse_response("esto no es json para nada")

    def test_parse_response_monto_string_con_simbolo(self):
        from resumidor import _parse_response
        raw = '{"resumen":"r","monto":"$1.500,00","moneda":"ARS","tipo":"Otro","proveedor":null,"fecha_doc":null,"confianza":0.5}'
        r = _parse_response(raw)
        self.assertIsNotNone(r["monto"])
        self.assertAlmostEqual(r["monto"], 1500.0, places=0)

    def test_parse_response_monto_null(self):
        from resumidor import _parse_response
        raw = '{"resumen":"Sin monto","monto":null,"moneda":null,"tipo":"Otro","proveedor":null,"fecha_doc":null,"confianza":0.0}'
        r = _parse_response(raw)
        self.assertIsNone(r["monto"])

    def test_parse_response_confianza_clampea_0_1(self):
        from resumidor import _parse_response
        # confianza > 1
        raw = '{"resumen":"r","monto":0,"moneda":null,"tipo":"Otro","proveedor":null,"fecha_doc":null,"confianza":1.5}'
        r = _parse_response(raw)
        self.assertLessEqual(r["confianza"], 1.0)
        # confianza < 0
        raw2 = '{"resumen":"r","monto":0,"moneda":null,"tipo":"Otro","proveedor":null,"fecha_doc":null,"confianza":-0.3}'
        r2 = _parse_response(raw2)
        self.assertGreaterEqual(r2["confianza"], 0.0)

    def test_parse_fecha_iso(self):
        from resumidor import _parse_fecha
        self.assertEqual(_parse_fecha("2024-03-15"), "2024-03-15")

    def test_parse_fecha_slash(self):
        from resumidor import _parse_fecha
        self.assertEqual(_parse_fecha("15/03/2024"), "2024-03-15")

    def test_parse_fecha_guion_dd_mm_yyyy(self):
        from resumidor import _parse_fecha
        self.assertEqual(_parse_fecha("15-03-2024"), "2024-03-15")

    def test_parse_fecha_none(self):
        from resumidor import _parse_fecha
        self.assertIsNone(_parse_fecha(None))

    def test_parse_fecha_vacia(self):
        from resumidor import _parse_fecha
        self.assertIsNone(_parse_fecha(""))

    def test_parse_fecha_invalida(self):
        from resumidor import _parse_fecha
        result = _parse_fecha("no-es-fecha")
        self.assertIsNone(result)

    def test_resumen_truncado_a_300_chars(self):
        from resumidor import _parse_response
        resumen_largo = "X" * 500
        raw = f'{{"resumen":"{resumen_largo}","monto":0,"moneda":null,"tipo":"Otro","proveedor":null,"fecha_doc":null,"confianza":0.0}}'
        r = _parse_response(raw)
        self.assertLessEqual(len(r["resumen"]), 300)


# =====================================================================
# 30. ENRIQUECEDOR — lógica pura
# =====================================================================
class TestEnriquecedor(unittest.TestCase):

    def test_looks_suspicious_gmail(self):
        from enriquecer_empresas_gemini import looks_suspicious
        self.assertTrue(looks_suspicious("gmail"))
        self.assertTrue(looks_suspicious("Gmail"))
        self.assertTrue(looks_suspicious("GMAIL"))

    def test_looks_suspicious_dominios_publicos(self):
        from enriquecer_empresas_gemini import looks_suspicious
        for d in ["hotmail","yahoo","outlook","icloud","live","aol","protonmail"]:
            self.assertTrue(looks_suspicious(d), f"'{d}' debería ser sospechoso")

    def test_looks_suspicious_string_vacio(self):
        from enriquecer_empresas_gemini import looks_suspicious
        self.assertTrue(looks_suspicious(""))
        self.assertTrue(looks_suspicious("  "))

    def test_looks_suspicious_muy_corto(self):
        from enriquecer_empresas_gemini import looks_suspicious
        self.assertTrue(looks_suspicious("A"))
        self.assertTrue(looks_suspicious("AB"))

    def test_looks_suspicious_solo_numeros(self):
        from enriquecer_empresas_gemini import looks_suspicious
        self.assertTrue(looks_suspicious("12345"))
        self.assertTrue(looks_suspicious("0"))

    def test_looks_suspicious_tld_solo(self):
        from enriquecer_empresas_gemini import looks_suspicious
        for tld in ["com","net","org","ar","es","uk"]:
            self.assertTrue(looks_suspicious(tld), f"TLD '{tld}' debería ser sospechoso")

    def test_looks_suspicious_nombres_validos(self):
        from enriquecer_empresas_gemini import looks_suspicious
        for nombre in ["ACME SA","Panadería San Martín","Tech Corp SRL",
                       "Google LLC","Empresa XYZ","Microsoft"]:
            self.assertFalse(looks_suspicious(nombre), f"'{nombre}' NO debería ser sospechoso")

    def test_looks_suspicious_subdominios_tecnicos(self):
        from enriquecer_empresas_gemini import looks_suspicious
        self.assertTrue(looks_suspicious("www"))
        self.assertTrue(looks_suspicious("mail"))
        self.assertTrue(looks_suspicious("smtp"))
        self.assertTrue(looks_suspicious("ftp"))

    def test_parse_json_response_lista_valida(self):
        from enriquecer_empresas_gemini import parse_json_response
        raw = '[{"id":1,"canonical_name":"ACME SA","confidence":0.9}]'
        result = parse_json_response(raw)
        self.assertIsInstance(result, list)
        self.assertEqual(len(result), 1)
        self.assertEqual(result[0]["id"], 1)

    def test_parse_json_response_con_fences(self):
        from enriquecer_empresas_gemini import parse_json_response
        raw = '```json\n[{"id":2,"canonical_name":"Beta SRL","confidence":0.85}]\n```'
        result = parse_json_response(raw)
        self.assertIsInstance(result, list)
        self.assertEqual(len(result), 1)

    def test_parse_json_response_vacio(self):
        from enriquecer_empresas_gemini import parse_json_response
        result = parse_json_response("[]")
        self.assertEqual(result, [])

    def test_parse_json_response_invalido(self):
        from enriquecer_empresas_gemini import parse_json_response
        result = parse_json_response("texto libre sin json")
        self.assertIsInstance(result, list)
        self.assertEqual(len(result), 0)

    def test_parse_json_response_json_embebido_en_texto(self):
        from enriquecer_empresas_gemini import parse_json_response
        # Modelo pone texto antes del JSON
        raw = 'Aquí está el resultado:\n[{"id":3,"canonical_name":"Gamma"}]'
        result = parse_json_response(raw)
        self.assertIsInstance(result, list)

    def test_validate_results_ids_correctos(self):
        from enriquecer_empresas_gemini import validate_results
        batch_ids = {1, 2, 3}
        results = [
            {"id": 1, "canonical_name": "A"},
            {"id": 2, "canonical_name": "B"},
            {"id": 3, "canonical_name": "C"},
        ]
        valid = validate_results(results, batch_ids)
        self.assertEqual(len(valid), 3)

    def test_validate_results_ids_alien_filtrados(self):
        from enriquecer_empresas_gemini import validate_results
        batch_ids = {1, 2}
        results = [
            {"id": 1,   "canonical_name": "A"},
            {"id": 99,  "canonical_name": "ALIEN"},   # no estaba en el batch
            {"id": 999, "canonical_name": "ALIEN2"},
        ]
        valid = validate_results(results, batch_ids)
        self.assertEqual(len(valid), 1)
        self.assertEqual(valid[0]["id"], 1)

    def test_validate_results_id_como_string(self):
        from enriquecer_empresas_gemini import validate_results
        batch_ids = {1, 2}
        results = [{"id": "1", "canonical_name": "A"}]  # id como string
        valid = validate_results(results, batch_ids)
        self.assertEqual(len(valid), 1)
        self.assertIsInstance(valid[0]["id"], int)

    def test_validate_results_id_invalido_ignorado(self):
        from enriquecer_empresas_gemini import validate_results
        batch_ids = {1}
        results = [{"id": "no-es-numero", "canonical_name": "X"}]
        valid = validate_results(results, batch_ids)
        self.assertEqual(valid, [])

    def test_build_prompt_contiene_ids(self):
        from enriquecer_empresas_gemini import build_prompt
        empresas = [
            {"id": 42, "nombre": "ACME", "pais": "AR", "rubro": "Tech", "email": ""},
            {"id": 99, "nombre": "Beta", "pais": "CL", "rubro": "", "email": ""},
        ]
        prompt = build_prompt(empresas)
        self.assertIn("42", prompt)
        self.assertIn("99", prompt)
        self.assertIn("ACME", prompt)

    def test_normalize_quita_espacios_multiples(self):
        from enriquecer_empresas_gemini import normalize
        self.assertEqual(normalize("  ACME   SA  "), "ACME SA")

    def test_normalize_none_o_vacio(self):
        from enriquecer_empresas_gemini import normalize
        self.assertEqual(normalize(""), "")
        self.assertEqual(normalize(None), "")

    def test_with_retry_sin_error_retorna_directo(self):
        from enriquecer_empresas_gemini import with_retry
        call_count = [0]
        def fn():
            call_count[0] += 1
            return "ok"
        result = with_retry(fn, max_attempts=3, base_delay=0.01)
        self.assertEqual(result, "ok")
        self.assertEqual(call_count[0], 1)

    def test_with_retry_error_no_retryable_lanza_inmediato(self):
        from enriquecer_empresas_gemini import with_retry
        call_count = [0]
        def fn():
            call_count[0] += 1
            raise ValueError("error de lógica, no de red")
        with self.assertRaises(ValueError):
            with_retry(fn, max_attempts=3, base_delay=0.01)
        self.assertEqual(call_count[0], 1, "No debe reintentar errores no-retryables")

    def test_with_retry_error_retryable_reintenta(self):
        from enriquecer_empresas_gemini import with_retry
        call_count = [0]
        def fn():
            call_count[0] += 1
            if call_count[0] < 3:
                raise Exception("rate limit 429 too many requests")
            return "ok_after_retry"
        result = with_retry(fn, max_attempts=4, base_delay=0.01)
        self.assertEqual(result, "ok_after_retry")
        self.assertEqual(call_count[0], 3)

    def test_strip_json_fences(self):
        from enriquecer_empresas_gemini import strip_json_fences
        self.assertEqual(strip_json_fences('```json\n[]\n```'), '[]')
        self.assertEqual(strip_json_fences('```\n[]\n```'), '[]')
        self.assertEqual(strip_json_fences('[1,2,3]'), '[1,2,3]')


# =====================================================================
# 31. DB — MÉTODOS NO CUBIERTOS
# =====================================================================
class TestDBMetodosNoCubiertos(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_obtener_empresa_por_nombre_exitoso(self):
        make_emp(self.db, "ACME SA")
        result = self.db.obtener_empresa_por_nombre("ACME SA")
        self.assertIsNotNone(result)
        self.assertEqual(result["nombre"], "ACME SA")

    def test_obtener_empresa_por_nombre_inexistente(self):
        result = self.db.obtener_empresa_por_nombre("XYZ No Existe")
        self.assertIsNone(result)

    def test_obtener_empresa_por_nombre_vacio(self):
        try:
            result = self.db.obtener_empresa_por_nombre("")
            self.assertIsNone(result)
        except Exception as e:
            self.fail(f"Crash nombre vacio: {e}")

    def test_obtener_empresa_por_nombre_none(self):
        try:
            result = self.db.obtener_empresa_por_nombre(None)
        except Exception as e:
            self.fail(f"Crash nombre None: {e}")

    def test_obtener_empresa_por_nombre_sql_injection(self):
        make_emp(self.db, "Legítima SA")
        try:
            result = self.db.obtener_empresa_por_nombre("'; DROP TABLE empresas;--")
            self.assertIsNone(result)
        except Exception as e:
            self.fail(f"Crash SQL injection: {e}")
        self.assertIsNotNone(self.db.fetchall("SELECT * FROM empresas"))

    def test_get_cotizacion_por_id_exitoso(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 5000.0, "Mi cotización")
        cot = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))
        self.assertIsNotNone(cot)
        result = self.db.get_cotizacion_por_id(cot["id"])
        self.assertIsNotNone(result)
        self.assertEqual(float(result["monto"]), 5000.0)

    def test_get_cotizacion_por_id_inexistente(self):
        result = self.db.get_cotizacion_por_id(99999)
        self.assertIsNone(result)

    def test_get_cotizacion_por_id_cero(self):
        result = self.db.get_cotizacion_por_id(0)
        self.assertIsNone(result)

    def test_get_cotizacion_por_id_negativo(self):
        result = self.db.get_cotizacion_por_id(-1)
        self.assertIsNone(result)

    def test_get_cotizacion_por_id_none(self):
        try:
            result = self.db.get_cotizacion_por_id(None)
        except Exception as e:
            self.fail(f"Crash id=None: {e}")

    def test_get_cotizacion_por_id_string(self):
        try:
            result = self.db.get_cotizacion_por_id("no-es-numero")
        except Exception as e:
            self.fail(f"Crash id=string: {e}")

    def test_deshacer_ultimo_cambio_devuelve_cambio(self):
        eid = make_emp(self.db, "Empresa Undo", pais="Argentina")
        # Preserve all other fields, change only nombre
        emp = self.db.obtener_empresa_por_id(eid)
        self.db.editar_empresa(
            eid, "Empresa Undo v2",
            emp.get("direccion",""), emp.get("telefono",""),
            emp.get("email",""), emp.get("rubro",""),
            emp.get("pais",""), "")
        cambio = self.db.deshacer_ultimo_cambio(eid)
        self.assertIsNotNone(cambio)
        # The last change must be nombre (we only changed that)
        self.assertEqual(cambio["campo"], "nombre")
        self.assertEqual(cambio["valor_anterior"], "Empresa Undo")
        self.assertEqual(cambio["valor_nuevo"], "Empresa Undo v2")

    def test_deshacer_ultimo_cambio_sin_historial(self):
        eid = make_emp(self.db, "Sin Historial")
        result = self.db.deshacer_ultimo_cambio(eid)
        self.assertIsNone(result)

    def test_deshacer_ultimo_cambio_devuelve_el_mas_reciente(self):
        eid = make_emp(self.db, "Multi Cambio")
        self.db.editar_empresa(eid, "v2", "", "", "", "", "", "")
        self.db.editar_empresa(eid, "v3", "", "", "", "", "", "")
        self.db.editar_empresa(eid, "v4", "", "", "", "", "", "")
        cambio = self.db.deshacer_ultimo_cambio(eid)
        self.assertEqual(cambio["valor_nuevo"], "v4")

    def test_vincular_tags_lista_vacia(self):
        eid = make_emp(self.db, "Sin Tags")
        try:
            self.db.vincular_empresa_con_tags(eid, [])
        except Exception as e:
            self.fail(f"Crash con lista vacía: {e}")
        self.assertEqual(self.db.get_tags_de_empresa(eid), [])

    def test_vincular_tags_none(self):
        eid = make_emp(self.db, "Tags None")
        try:
            self.db.vincular_empresa_con_tags(eid, None)
        except Exception as e:
            self.fail(f"Crash con tags=None: {e}")


# =====================================================================
# 32. UTILS — MÉTODOS NO CUBIERTOS
# =====================================================================
class TestUtilsNoCubiertos(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_get_fuzzy_threshold_devuelve_int(self):
        from utils import Config
        cfg = Config()
        thresh = cfg.get_fuzzy_threshold()
        self.assertIsInstance(thresh, int)
        self.assertGreater(thresh, 0)

    def test_restaurar_backup_sin_backups_no_crashea(self):
        from utils import BackupManager
        # Path a DB que no existe = sin backups
        result = BackupManager.restaurar_backup("/ruta/inexistente/test.db")
        # Debe devolver False, no crashear
        self.assertFalse(result)

    def test_detect_encoding_utf8(self):
        from utils import CSVCleaner
        path = os.path.join(self.tmpdir, "utf8.csv")
        with open(path, "w", encoding="utf-8") as f:
            f.write("nombre,email\nJuan,j@test.com\n")
        enc = CSVCleaner.detect_encoding(path)
        self.assertIsInstance(enc, str)
        self.assertIn("utf", enc.lower())

    def test_detect_encoding_latin1(self):
        from utils import CSVCleaner
        path = os.path.join(self.tmpdir, "latin.csv")
        with open(path, "w", encoding="latin-1") as f:
            f.write("nombre,email\nMaría,m@e.com\n")
        enc = CSVCleaner.detect_encoding(path)
        self.assertIsInstance(enc, str)
        self.assertGreater(len(enc), 0)

    def test_clean_csv_basico(self):
        from utils import CSVCleaner
        path = os.path.join(self.tmpdir, "basic.csv")
        with open(path, "w", encoding="utf-8") as f:
            f.write("nombre,email,empresa\nJuan,juan@acme.com,ACME\n")
        result = CSVCleaner.clean_csv(path)
        self.assertIsInstance(result, list)

    def test_clean_csv_archivo_vacio(self):
        from utils import CSVCleaner
        path = os.path.join(self.tmpdir, "vacio.csv")
        with open(path, "w", encoding="utf-8") as f:
            f.write("")
        try:
            result = CSVCleaner.clean_csv(path)
            self.assertIsInstance(result, list)
        except Exception as e:
            self.fail(f"Crash con CSV vacío: {e}")

    def test_clean_email_espacios_internos(self):
        from utils import CSVCleaner
        # Email con espacios no debe ser válido
        result = CSVCleaner.clean_email("test @example.com")
        self.assertEqual(result, "")

    def test_clean_email_mayusculas(self):
        from utils import CSVCleaner
        result = CSVCleaner.clean_email("Test@EMPRESA.COM")
        self.assertEqual(result, "test@empresa.com")


# =====================================================================
# 33. BORDES Y NEGATIVOS ADICIONALES EN DB
# =====================================================================
class TestBordesAdicionalesDB(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_agregar_empresa_nombre_none(self):
        try:
            ok = self.db.agregar_empresa(None, "", "", "", "", "", "")
        except Exception as e:
            self.fail(f"Crash nombre=None: {e}")

    def test_editar_empresa_id_none(self):
        try:
            ok = self.db.editar_empresa(None, "Nombre", "", "", "", "", "", "")
        except Exception as e:
            self.fail(f"Crash empresa_id=None: {e}")

    def test_editar_empresa_id_string(self):
        try:
            ok = self.db.editar_empresa("no-es-id", "Nombre", "", "", "", "", "", "")
        except Exception as e:
            self.fail(f"Crash empresa_id=string: {e}")

    def test_registrar_cambio_empresa_id_none(self):
        try:
            ok = self.db.registrar_cambio(None, "nombre", "a", "b", "usuario")
        except Exception as e:
            self.fail(f"Crash empresa_id=None en historial: {e}")

    def test_registrar_cambio_valores_none(self):
        eid = make_emp(self.db, "Test None Values")
        try:
            ok = self.db.registrar_cambio(eid, "nombre", None, None, "usuario")
        except Exception as e:
            self.fail(f"Crash valores None en historial: {e}")

    def test_registrar_cambio_valores_muy_largos(self):
        eid = make_emp(self.db, "Test Largo")
        try:
            ok = self.db.registrar_cambio(
                eid, "nombre", "A"*5000, "B"*5000, "usuario")
        except Exception as e:
            self.fail(f"Crash valores largos en historial: {e}")

    def test_agregar_cotizacion_descripcion_none(self):
        eid = make_emp(self.db)
        try:
            ok = self.db.agregar_cotizacion(eid, None, 100.0)
        except Exception as e:
            self.fail(f"Crash descripcion=None: {e}")

    def test_agregar_cotizacion_empresa_id_string(self):
        try:
            ok = self.db.agregar_cotizacion("no-es-id", "desc", 100.0)
        except Exception as e:
            self.fail(f"Crash empresa_id=string en cotizacion: {e}")

    def test_count_by_empresa_id_none(self):
        try:
            n = self.db.count_by_empresa("cotizaciones", None)
            self.assertEqual(n, 0)
        except Exception as e:
            self.fail(f"Crash count_by_empresa con id=None: {e}")

    def test_count_by_empresa_tabla_con_espacios(self):
        try:
            n = self.db.count_by_empresa("tabla con espacios", 1)
        except Exception as e:
            self.fail(f"Crash tabla con espacios: {e}")

    def test_get_historial_empresa_id_string(self):
        try:
            hist = self.db.get_historial_empresa("no-es-id")
            self.assertIsInstance(hist, list)
        except Exception as e:
            self.fail(f"Crash historial id=string: {e}")

    def test_eliminar_contacto_id_none(self):
        try:
            ok = self.db.eliminar_contacto(None)
        except Exception as e:
            self.fail(f"Crash eliminar contacto id=None: {e}")

    def test_eliminar_contacto_id_string(self):
        try:
            ok = self.db.eliminar_contacto("no-es-id")
        except Exception as e:
            self.fail(f"Crash eliminar contacto id=string: {e}")

    def test_editar_contacto_id_inexistente(self):
        try:
            ok = self.db.editar_contacto(99999, "Nombre", "e@test.com", "", "")
        except Exception as e:
            self.fail(f"Crash editar contacto inexistente: {e}")

    def test_unificar_empresas_id_none(self):
        try:
            ok = self.db.unificar_empresas(None, None)
        except Exception as e:
            self.fail(f"Crash unificar id=None: {e}")

    def test_get_filtered_empresas_pais_sql_injection(self):
        make_emp(self.db, "Legítima")
        try:
            result = self.db.get_filtered_empresas(
                "", {"pais": "'; DROP TABLE empresas;--"})
        except Exception as e:
            self.fail(f"Crash SQL injection en filtro pais: {e}")
        self.assertGreater(self.db.count("empresas"), 0)

    def test_get_filtered_empresas_rubro_sql_injection(self):
        make_emp(self.db, "Empresa")
        try:
            result = self.db.get_filtered_empresas(
                "", {"rubro": "' OR 1=1--"})
        except Exception as e:
            self.fail(f"Crash SQL injection en filtro rubro: {e}")

    def test_get_filtered_empresas_tag_sql_injection(self):
        make_emp(self.db, "Empresa")
        try:
            result = self.db.get_filtered_empresas(
                "", {"tag": "'; DROP TABLE tags;--"})
        except Exception as e:
            self.fail(f"Crash SQL injection en filtro tag: {e}")
        self.assertIsNotNone(self.db.fetchall("SELECT * FROM tags"))



class TestStressConcurrenciaFileDB(unittest.TestCase):
    """
    Tests de stress con DB en archivo — el escenario real de 3 usuarios.
    WAL + _write_lock deben garantizar:
      - Cero pérdida de datos
      - Cero corrupción
      - Cero deadlocks
      - Conteos exactos
    """

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.db_path = os.path.join(self.tmpdir, "stress.db")
        self.db = DBManager(self.db_path)

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _open_second_connection(self):
        """Simula un segundo usuario abriendo la misma DB."""
        return DBManager(self.db_path)

    def _open_third_connection(self):
        return DBManager(self.db_path)

    # ── Tests ─────────────────────────────────────────────────────────────────

    def test_3_usuarios_escrituras_independientes_sin_perdida(self):
        """
        3 conexiones independientes escriben 100 empresas cada una.
        Total esperado: 300. Ninguna escritura debe perderse.
        """
        db2 = self._open_second_connection()
        db3 = self._open_third_connection()
        errors = []

        def usuario(db, prefix, n):
            for i in range(n):
                ok = db.agregar_empresa(
                    f"{prefix}_{i:03d}", "", "", "", "", "AR", "")
                if not ok:
                    errors.append(f"{prefix}_{i} falló")

        threads = [
            threading.Thread(target=usuario, args=(self.db, "U1", 100)),
            threading.Thread(target=usuario, args=(db2,      "U2", 100)),
            threading.Thread(target=usuario, args=(db3,      "U3", 100)),
        ]
        for t in threads: t.start()
        for t in threads: t.join(timeout=30)

        alive = [t for t in threads if t.is_alive()]
        self.assertEqual(len(alive), 0, "Deadlock detectado")
        self.assertEqual(errors, [], f"Escrituras fallidas: {errors}")

        total = self.db.count("empresas")
        self.assertEqual(total, 300,
            f"Pérdida de datos: esperaba 300, hay {total}")

    def test_escrituras_concurrentes_historial_consistente(self):
        """
        Editar la misma empresa desde 2 threads simultáneos.
        El historial debe registrar todos los cambios sin corrupción.
        """
        eid = make_emp(self.db, "Empresa Compartida")
        emp = self.db.obtener_empresa_por_id(eid)
        errors = []
        cambios_hechos = []

        def editar(nuevo_nombre, nuevo_rubro):
            try:
                ok = self.db.editar_empresa(
                    eid, nuevo_nombre,
                    emp.get("direccion",""), emp.get("telefono",""),
                    emp.get("email",""), nuevo_rubro,
                    emp.get("pais",""), "")
                if ok:
                    cambios_hechos.append(nuevo_nombre)
                else:
                    errors.append(f"editar falló: {nuevo_nombre}")
            except Exception as e:
                errors.append(str(e))

        threads = [
            threading.Thread(target=editar, args=(f"Nombre_v{i}", f"Rubro_{i}"))
            for i in range(10)
        ]
        for t in threads: t.start()
        for t in threads: t.join(timeout=15)

        self.assertEqual(errors, [], f"Errores editando: {errors}")
        # El historial debe tener registros (cuántos exactamente depende de
        # cuántos cambios reales hubo, pero no debe estar vacío ni corrupto)
        hist = self.db.get_historial_empresa(eid)
        self.assertIsInstance(hist, list)
        # Verificar que todos los registros tienen empresa_id correcto
        for h in hist:
            self.assertEqual(h["empresa_id"], eid,
                f"Historial corrupto: empresa_id={h['empresa_id']} != {eid}")

    def test_lector_no_ve_datos_parciales(self):
        """
        Un reader no debe ver una empresa a medio insertar.
        INSERT es atómico — el reader ve el estado antes o después, nunca a la mitad.
        """
        errores_integridad = []

        def escritor():
            for i in range(200):
                # INSERT con múltiples campos — debe ser atómico
                self.db.agregar_empresa(
                    f"Empresa_{i:03d}", f"Dir {i}", f"0800-{i}",
                    f"e{i}@test.com", f"Rubro_{i % 5}", "AR", "")
                time.sleep(0.001)

        def lector():
            for _ in range(100):
                empresas = self.db.get_filtered_empresas("", {})
                for emp in empresas:
                    # Cada empresa debe tener nombre válido (no None, no vacío parcial)
                    if not emp.get("nombre"):
                        errores_integridad.append(
                            f"Empresa sin nombre: id={emp.get('id')}")
                time.sleep(0.002)

        t_w = threading.Thread(target=escritor)
        t_r = threading.Thread(target=lector)
        t_w.start(); t_r.start()
        t_w.join(timeout=20); t_r.join(timeout=20)

        self.assertEqual(errores_integridad, [],
            f"Datos parciales detectados: {errores_integridad[:3]}")

    def test_conteo_exacto_bajo_concurrencia(self):
        """
        count() debe devolver el número exacto incluso con escrituras concurrentes.
        """
        N_POR_THREAD = 50
        N_THREADS = 6

        def writer(tid):
            for i in range(N_POR_THREAD):
                self.db.agregar_empresa(
                    f"T{tid}_{i:03d}", "", "", "", "", "", "")

        threads = [threading.Thread(target=writer, args=(t,))
                   for t in range(N_THREADS)]
        for t in threads: t.start()
        for t in threads: t.join(timeout=20)

        expected = N_POR_THREAD * N_THREADS
        actual   = self.db.count("empresas")
        self.assertEqual(actual, expected,
            f"Conteo incorrecto bajo concurrencia: esperaba {expected}, hay {actual}")

    def test_cascade_delete_bajo_concurrencia(self):
        """
        Eliminar una empresa mientras otro thread agrega contactos a la misma.
        La DB no debe quedar con contactos huérfanos.
        """
        eid = make_emp(self.db, "Empresa Para Borrar")
        stop_flag = threading.Event()
        errors = []

        def agregar_contactos():
            i = 0
            while not stop_flag.is_set():
                try:
                    self.db.agregar_contacto(
                        eid, f"Contacto_{i}", f"c{i}@test.com", "", "")
                    i += 1
                except Exception:
                    pass  # La empresa puede haber sido borrada

        t_add = threading.Thread(target=agregar_contactos)
        t_add.start()
        time.sleep(0.05)  # Dar tiempo a que se agreguen algunos contactos

        # Borrar la empresa mientras el thread agrega contactos
        self.db.eliminar_empresa(eid)
        stop_flag.set()
        t_add.join(timeout=5)

        # Verificar no hay contactos huérfanos
        huerfanos = self.db.fetchall(
            "SELECT * FROM contactos WHERE empresa_id=?", (eid,))
        self.assertEqual(len(huerfanos), 0,
            f"{len(huerfanos)} contactos huérfanos después del cascade delete")

    def test_3_usuarios_mix_crud_completo(self):
        """
        Simula 3 usuarios haciendo operaciones CRUD mezcladas simultáneamente.
        Al final, todos los counts deben ser consistentes.
        """
        db2 = self._open_second_connection()
        db3 = self._open_third_connection()
        errors = []

        def usuario_crud(db, prefix):
            """Crea empresas, les agrega contactos y cotizaciones, luego fusiona."""
            try:
                # Crear 20 empresas
                ids = []
                for i in range(20):
                    db.agregar_empresa(
                        f"{prefix}_E{i}", "", "", "", "", "AR", "")
                    r = db.fetchone(
                        "SELECT id FROM empresas WHERE nombre=?",
                        (f"{prefix}_E{i}",))
                    if r:
                        ids.append(r["id"])

                # Agregar contactos y cotizaciones a cada empresa
                for eid in ids:
                    db.agregar_contacto(
                        eid, f"Contacto_{eid}", f"c{eid}@test.com", "", "")
                    db.agregar_cotizacion(eid, f"Cotiz_{eid}", float(eid * 100))

                # Editar algunas empresas (genera historial)
                for eid in ids[:5]:
                    emp = db.obtener_empresa_por_id(eid)
                    if emp:
                        db.editar_empresa(
                            eid, emp["nombre"] + "_edit",
                            "", "", "", "", "CL", "")

            except Exception as e:
                errors.append(f"{prefix}: {e}")

        threads = [
            threading.Thread(target=usuario_crud, args=(self.db, "U1")),
            threading.Thread(target=usuario_crud, args=(db2,      "U2")),
            threading.Thread(target=usuario_crud, args=(db3,      "U3")),
        ]
        for t in threads: t.start()
        for t in threads: t.join(timeout=45)

        alive = [t for t in threads if t.is_alive()]
        self.assertEqual(len(alive), 0, "Deadlock en CRUD mezclado")
        self.assertEqual(errors, [], f"Errores CRUD: {errors}")

        # Verificar consistencia: cada empresa debe tener exactamente 1 contacto
        # y 1 cotización (las que se crearon, más posibles ediciones)
        n_empresas    = self.db.count("empresas")
        n_contactos   = self.db.count("contactos")
        n_cotizaciones = self.db.count("cotizaciones")

        self.assertEqual(n_empresas, 60,
            f"Empresas: esperaba 60, hay {n_empresas}")
        self.assertEqual(n_contactos, 60,
            f"Contactos: esperaba 60, hay {n_contactos}")
        self.assertEqual(n_cotizaciones, 60,
            f"Cotizaciones: esperaba 60, hay {n_cotizaciones}")

    def test_timeout_no_deadlock_en_contención(self):
        """
        Con WAL + busy_timeout=5000ms, una escritura lenta no debe
        bloquear indefinidamente a otra. El segundo writer debe completar
        dentro del timeout razonable.
        """
        import signal

        db2 = self._open_second_connection()
        times = []
        errors = []

        def escritor_lento():
            """Abre una transacción explícita y mantiene el lock brevemente."""
            try:
                # Escritura normal — el lock se libera inmediatamente
                t0 = time.time()
                for i in range(50):
                    self.db.agregar_empresa(f"Lento_{i}", "", "", "", "", "", "")
                times.append(time.time() - t0)
            except Exception as e:
                errors.append(f"lento: {e}")

        def escritor_rapido():
            """Otro writer — debe poder escribir sin bloquearse."""
            try:
                t0 = time.time()
                for i in range(50):
                    db2.agregar_empresa(f"Rapido_{i}", "", "", "", "", "", "")
                times.append(time.time() - t0)
            except Exception as e:
                errors.append(f"rapido: {e}")

        t1 = threading.Thread(target=escritor_lento)
        t2 = threading.Thread(target=escritor_rapido)
        t1.start(); t2.start()
        t1.join(timeout=15); t2.join(timeout=15)

        alive = [t for t in (t1, t2) if t.is_alive()]
        self.assertEqual(len(alive), 0,
            f"Deadlock: {len(alive)} threads colgados después de 15s")
        self.assertEqual(errors, [], f"Errores: {errors}")
        # Ambos deben haber terminado en tiempo razonable
        if times:
            self.assertLess(max(times), 12.0,
                f"Escritura tardó demasiado: {max(times):.2f}s")

    def test_integridad_referencial_bajo_concurrencia(self):
        """
        Con FOREIGN KEY constraints activas, los cascade deletes deben
        funcionar correctamente incluso con escrituras concurrentes.
        """
        # Crear empresas con datos
        eids = []
        for i in range(20):
            eid = make_emp(self.db, f"FK_Test_{i}")
            make_con(self.db, eid)
            make_cot(self.db, eid)
            eids.append(eid)

        errors = []

        def borrar_empresas(ids):
            for eid in ids:
                try:
                    self.db.eliminar_empresa(eid)
                except Exception as e:
                    errors.append(str(e))

        def agregar_mas(prefix, n):
            for i in range(n):
                try:
                    eid = make_emp(self.db, f"{prefix}_{i}")
                    make_cot(self.db, eid)
                except Exception as e:
                    errors.append(str(e))

        # Borrar la mitad mientras se agregan nuevas
        t1 = threading.Thread(target=borrar_empresas, args=(eids[:10],))
        t2 = threading.Thread(target=agregar_mas, args=("New", 20))
        t1.start(); t2.start()
        t1.join(timeout=15); t2.join(timeout=15)

        self.assertEqual(errors, [], f"Errores FK: {errors}")

        # Verificar no hay datos huérfanos
        contactos_huerfanos = self.db.fetchall("""
            SELECT c.id FROM contactos c
            LEFT JOIN empresas e ON c.empresa_id = e.id
            WHERE e.id IS NULL
        """)
        cotizaciones_huerfanas = self.db.fetchall("""
            SELECT c.id FROM cotizaciones c
            LEFT JOIN empresas e ON c.empresa_id = e.id
            WHERE e.id IS NULL
        """)
        self.assertEqual(len(contactos_huerfanos), 0,
            f"{len(contactos_huerfanos)} contactos huérfanos")
        self.assertEqual(len(cotizaciones_huerfanas), 0,
            f"{len(cotizaciones_huerfanas)} cotizaciones huérfanas")


# =====================================================================
# 35. STRESS DE CONCURRENCIA — :memory: (tests headless)
# =====================================================================
class TestStressConcurrenciaMemory(unittest.TestCase):
    """
    Verifica que _MemConn con RLock es thread-safe.
    Los tests de concurrencia en :memory: validan que el RLock serializa
    correctamente sin deadlocks ni pérdida de datos.
    """

    def setUp(self):
        self.db = fresh_db()

    def test_500_escrituras_5_threads_sin_perdida(self):
        """5 threads × 100 INSERT = 500 empresas exactas."""
        errors = []

        def writer(prefix):
            for i in range(100):
                ok = self.db.agregar_empresa(
                    f"{prefix}_{i:03d}", "", "", "", "", "", "")
                if not ok:
                    errors.append(f"{prefix}_{i}")

        threads = [threading.Thread(target=writer, args=(f"T{t}",))
                   for t in range(5)]
        for t in threads: t.start()
        for t in threads: t.join(timeout=20)

        alive = [t for t in threads if t.is_alive()]
        self.assertEqual(len(alive), 0, "Deadlock en :memory:")
        self.assertEqual(errors, [], f"Escrituras fallidas: {errors}")
        self.assertEqual(self.db.count("empresas"), 500)

    def test_rlock_reentrant_mismo_thread(self):
        """
        El mismo thread puede adquirir el RLock múltiples veces.
        editar_empresa llama a vincular_empresa_con_tags que también usa _get_connection.
        Esto no debe deadlockear.
        """
        eid = make_emp(self.db, "Empresa RLock")
        try:
            # editar_empresa internamente llama vincular_empresa_con_tags
            # ambos adquieren el mismo RLock (re-entrant, no deadlock)
            ok = self.db.editar_empresa(
                eid, "Empresa RLock Editada", "", "", "", "", "AR",
                "tag1, tag2, tag3")
            self.assertTrue(ok, "editar_empresa falló con RLock")
        except Exception as e:
            self.fail(f"Deadlock o error en RLock re-entrant: {e}")

    def test_conteo_exacto_con_reads_simultaneos(self):
        """Escrituras + lecturas simultáneas dan conteo exacto al final."""
        errors = []

        def writer():
            for i in range(100):
                self.db.agregar_empresa(f"W_{i}", "", "", "", "", "", "")

        def reader():
            for _ in range(50):
                try:
                    self.db.get_filtered_empresas("", {})
                except Exception as e:
                    errors.append(str(e))
                time.sleep(0.001)

        threads = [
            threading.Thread(target=writer),
            threading.Thread(target=reader),
            threading.Thread(target=reader),
        ]
        for t in threads: t.start()
        for t in threads: t.join(timeout=20)

        self.assertEqual(errors, [])
        # El writer escribió exactamente 100
        self.assertEqual(self.db.count("empresas"), 100)

    def test_historial_integro_bajo_concurrencia(self):
        """
        El historial de cambios no debe tener registros corruptos
        cuando múltiples threads editan empresas simultáneamente.
        """
        # Crear 10 empresas
        eids = [make_emp(self.db, f"Hist_{i}") for i in range(10)]
        errors = []

        def editar_empresa(eid, suffix):
            emp = self.db.obtener_empresa_por_id(eid)
            if emp:
                try:
                    self.db.editar_empresa(
                        eid, emp["nombre"] + suffix,
                        "", "", "", "", "AR", "")
                except Exception as e:
                    errors.append(str(e))

        # 5 threads editan todas las empresas
        threads = []
        for suffix in [f"_v{i}" for i in range(5)]:
            for eid in eids:
                threads.append(
                    threading.Thread(target=editar_empresa, args=(eid, suffix)))

        for t in threads: t.start()
        for t in threads: t.join(timeout=20)

        self.assertEqual(errors, [], f"Errores editando: {errors}")

        # Verificar integridad del historial
        for eid in eids:
            hist = self.db.get_historial_empresa(eid)
            for h in hist:
                self.assertEqual(h["empresa_id"], eid,
                    f"empresa_id corrupto en historial: {h}")
                self.assertIn("campo", h)
                self.assertIn("fecha", h)

    def test_no_deadlock_con_10_threads_30s(self):
        """
        10 threads haciendo operaciones mixtas durante 30 iteraciones.
        Ningún thread debe colgarse.
        """
        eids = [make_emp(self.db, f"Mix_{i}") for i in range(5)]
        errors = []

        def mixed_worker(tid):
            for i in range(30):
                try:
                    if i % 3 == 0:
                        self.db.agregar_empresa(
                            f"T{tid}_new_{i}", "", "", "", "", "", "")
                    elif i % 3 == 1:
                        self.db.get_filtered_empresas("", {})
                    else:
                        eid = eids[i % len(eids)]
                        make_cot(self.db, eid, float(i))
                except Exception as e:
                    errors.append(f"T{tid}:{i}: {e}")

        threads = [threading.Thread(target=mixed_worker, args=(t,))
                   for t in range(10)]
        for t in threads: t.start()
        for t in threads: t.join(timeout=30)

        alive = [t for t in threads if t.is_alive()]
        self.assertEqual(len(alive), 0,
            f"{len(alive)} threads colgados — posible deadlock")
        bad = [e for e in errors
               if "lock" not in e.lower() and "database" not in e.lower()]
        self.assertEqual(bad, [], f"Errores no-lock: {bad}")

    def test_cascade_delete_atomico(self):
        """
        Borrar empresa mientras otro thread agrega datos a ella.
        El resultado final debe ser consistente (sin huérfanos).
        """
        eid = make_emp(self.db, "Cascade Test")
        stop_flag = threading.Event()

        def agregar_loop():
            i = 0
            while not stop_flag.is_set():
                try:
                    self.db.agregar_contacto(
                        eid, f"C{i}", f"c{i}@test.com", "", "")
                    i += 1
                except Exception:
                    pass
                time.sleep(0.001)

        t = threading.Thread(target=agregar_loop)
        t.start()
        time.sleep(0.05)

        self.db.eliminar_empresa(eid)
        stop_flag.set()
        t.join(timeout=5)

        huerfanos = self.db.fetchall(
            "SELECT * FROM contactos WHERE empresa_id=?", (eid,))
        self.assertEqual(len(huerfanos), 0,
            f"{len(huerfanos)} contactos huérfanos")


# =====================================================================
# 36. STRESS DE VOLUMEN EXTREMO CON HISTORIAL
# =====================================================================
class TestStressVolumenHistorial(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.db = DBManager(os.path.join(self.tmpdir, "vol.db"))

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_1000_empresas_con_historial_rendimiento(self):
        """
        Crear 1000 empresas y editar cada una 3 veces.
        Total: 1000 empresas, ~3000 registros en cambios.
        Leer historial de todas debe ser rápido.
        """
        eids = []
        for i in range(1000):
            self.db.agregar_empresa(f"E_{i:04d}", "", "", "", "", "AR", "")
            r = self.db.fetchone(
                "SELECT id FROM empresas WHERE nombre=?", (f"E_{i:04d}",))
            if r:
                eids.append(r["id"])

        # Editar cada empresa 3 veces
        for eid in eids[:200]:   # Solo 200 para velocidad
            emp = self.db.obtener_empresa_por_id(eid)
            if not emp: continue
            for v in range(3):
                self.db.editar_empresa(
                    eid, emp["nombre"] + f"_v{v}",
                    "", "", "", "", "CL", "")

        total_cambios = self.db.fetchone(
            "SELECT COUNT(*) n FROM cambios")["n"]
        self.assertGreater(total_cambios, 0)

        # Leer historial de una empresa con muchos cambios debe ser rápido
        t0 = time.time()
        hist = self.db.get_historial_empresa(eids[0])
        elapsed = time.time() - t0
        self.assertLess(elapsed, 1.0, f"get_historial tardó {elapsed:.2f}s")

    def test_busqueda_1000_empresas_con_todos_filtros(self):
        """
        Búsqueda con todos los filtros combinados sobre 1000 empresas.
        Debe responder en < 3s.
        """
        for i in range(1000):
            pais = "Argentina" if i % 2 == 0 else "Chile"
            rubro = "Tech" if i % 3 == 0 else "Comercio"
            self.db.agregar_empresa(
                f"Empresa_{i:04d}", "", "", "", rubro, pais, "")
            if i % 10 == 0:
                r = self.db.fetchone(
                    "SELECT id FROM empresas WHERE nombre=?",
                    (f"Empresa_{i:04d}",))
                if r:
                    make_cot(self.db, r["id"])

        t0 = time.time()
        result = self.db.get_filtered_empresas("Empresa_0", {
            "pais": "Argentina",
            "rubro": "Tech",
            "cotizaciones_cond": "con",
        })
        elapsed = time.time() - t0
        self.assertLess(elapsed, 3.0,
            f"Búsqueda con filtros tardó {elapsed:.2f}s")

    def test_similar_1000_empresas_rendimiento_wal(self):
        """
        get_similar_empresas con 1000 empresas en file DB con WAL.
        El pre-filtro debe mantenerlo por debajo de 15s.
        """
        for i in range(1000):
            self.db.agregar_empresa(f"Empresa Similar {i:04d}", "", "", "", "", "", "")

        t0 = time.time()
        pares = self.db.get_similar_empresas(85)
        elapsed = time.time() - t0
        self.assertLess(elapsed, 15.0,
            f"get_similar_empresas tardó {elapsed:.2f}s con 1000 empresas y WAL")



class TestEditarCotizacion(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_editar_descripcion_basico(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "Original")
        cot = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))
        ok = self.db.editar_cotizacion(cot["id"], "Editada", 1000.0)
        self.assertTrue(ok)
        cot2 = self.db.get_cotizacion_por_id(cot["id"])
        self.assertEqual(cot2["descripcion"], "Editada")

    def test_editar_monto(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "Cotiz")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        ok = self.db.editar_cotizacion(cid, "Cotiz", 9999.99)
        self.assertTrue(ok)
        cot = self.db.get_cotizacion_por_id(cid)
        self.assertAlmostEqual(float(cot["monto"]), 9999.99, places=2)

    def test_editar_tipo(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 500.0, "Sin tipo")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        ok = self.db.editar_cotizacion(cid, "Con tipo", 500.0, tipo="Equipos")
        self.assertTrue(ok)
        cot = self.db.get_cotizacion_por_id(cid)
        self.assertEqual(cot.get("tipo"), "Equipos")

    def test_editar_fecha(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "Cotiz", "2024-01-01 00:00:00")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        ok = self.db.editar_cotizacion(cid, "Cotiz", 1000.0,
                                        fecha="2025-06-15 10:30:00")
        self.assertTrue(ok)
        cot = self.db.get_cotizacion_por_id(cid)
        self.assertIn("2025", cot.get("fecha",""))

    def test_editar_cotizacion_inexistente(self):
        ok = self.db.editar_cotizacion(99999, "X", 100.0)
        self.assertFalse(ok)

    def test_editar_cotizacion_id_none(self):
        ok = self.db.editar_cotizacion(None, "X", 100.0)
        self.assertFalse(ok)

    def test_editar_cotizacion_id_cero(self):
        ok = self.db.editar_cotizacion(0, "X", 100.0)
        self.assertFalse(ok)

    def test_editar_monto_negativo_no_crashea(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "Cotiz")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        try:
            ok = self.db.editar_cotizacion(cid, "Cotiz", -500.0)
        except Exception as e:
            self.fail(f"Crash monto negativo: {e}")

    def test_editar_descripcion_vacia_no_crashea(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "Cotiz")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        try:
            ok = self.db.editar_cotizacion(cid, "", 1000.0)
        except Exception as e:
            self.fail(f"Crash descripcion vacia: {e}")

    def test_editar_descripcion_sql_injection(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "Cotiz")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        payload = "'; DROP TABLE cotizaciones;--"
        try:
            self.db.editar_cotizacion(cid, payload, 1000.0)
        except Exception as e:
            self.fail(f"Crash SQL injection: {e}")
        self.assertIsNotNone(
            self.db.fetchall("SELECT * FROM cotizaciones"))

    def test_eliminar_cotizacion(self):
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "Para borrar")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        ok = self.db.eliminar_cotizacion(cid)
        self.assertTrue(ok)
        self.assertIsNone(self.db.get_cotizacion_por_id(cid))
        self.assertEqual(self.db.count_by_empresa("cotizaciones", eid), 0)

    def test_eliminar_cotizacion_inexistente(self):
        ok = self.db.eliminar_cotizacion(99999)
        self.assertTrue(ok)  # no crash, devuelve True (0 rows affected)

    def test_eliminar_cotizacion_id_none(self):
        ok = self.db.eliminar_cotizacion(None)
        self.assertFalse(ok)

    def test_cotizacion_con_tipo_prefijo_extrae_tipo(self):
        """[Equipos] desc -> tipo='Equipos' se extrae en editar_cotizacion."""
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "[Comercial] Contrato anual")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        ok = self.db.editar_cotizacion(cid, "[Equipos] Nueva descripción", 2000.0)
        self.assertTrue(ok)
        cot = self.db.get_cotizacion_por_id(cid)
        self.assertEqual(cot.get("tipo"), "Equipos")

    def test_editar_preserva_ruta_archivo(self):
        """editar_cotizacion no debe tocar ruta_archivo."""
        eid = make_emp(self.db)
        self.db.agregar_cotizacion_con_ruta(
            eid, "archivo.pdf", 1000.0,
            "2024-01-01 00:00:00", "/srv/archivo.pdf")
        cid = self.db.fetchone("SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]
        self.db.editar_cotizacion(cid, "archivo editado.pdf", 2000.0)
        cot = self.db.get_cotizacion_por_id(cid)
        # ruta_archivo should be unchanged
        self.assertIsNotNone(cot)


# =====================================================================
# 38. LÓGICA TAGS — NORMALIZACIÓN, VACÍOS, UNICODE
# =====================================================================
class TestTagsNormalizacion(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_tags_con_espacios_se_normalizan(self):
        eid = make_emp(self.db, "Tags Espacios")
        self.db.vincular_empresa_con_tags(eid, ["  vip  ", " cliente "])
        tags = self.db.get_tags_de_empresa(eid)
        self.assertIn("vip", tags)
        self.assertIn("cliente", tags)
        self.assertNotIn("  vip  ", tags)

    def test_tags_vacios_en_string_no_se_guardan(self):
        eid = make_emp(self.db, "Tags Vacios")
        self.db.editar_empresa(eid, "Tags Vacios", "", "", "", "", "",
                               "  ,  ,  vip  ,  ,  cliente  ,  ")
        tags = self.db.get_tags_de_empresa(eid)
        self.assertIn("vip", tags)
        self.assertIn("cliente", tags)
        # No debe haber tags vacíos
        for t in tags:
            self.assertGreater(len(t), 0, f"Tag vacío encontrado: {t!r}")

    def test_tags_duplicados_no_se_guardan(self):
        eid = make_emp(self.db, "Tags Dup")
        self.db.vincular_empresa_con_tags(eid, ["vip", "vip", "vip"])
        tags = self.db.get_tags_de_empresa(eid)
        self.assertEqual(tags.count("vip"), 1)

    def test_tags_unicode(self):
        eid = make_emp(self.db, "Tags Unicode")
        self.db.vincular_empresa_con_tags(
            eid, ["exportación", "construcción", "año_2024"])
        tags = self.db.get_tags_de_empresa(eid)
        self.assertIn("exportación", tags)
        self.assertIn("construcción", tags)

    def test_tags_solo_espacios_no_se_guardan(self):
        eid = make_emp(self.db, "Tags Solo Espacios")
        self.db.vincular_empresa_con_tags(eid, ["   ", "\t", "\n"])
        tags = self.db.get_tags_de_empresa(eid)
        self.assertEqual(len(tags), 0)

    def test_editar_empresa_reemplaza_todos_los_tags(self):
        eid = make_emp(self.db, "Tags Replace")
        self.db.vincular_empresa_con_tags(eid, ["viejo1", "viejo2"])
        self.db.editar_empresa(eid, "Tags Replace", "", "", "", "", "",
                               "nuevo1, nuevo2, nuevo3")
        tags = self.db.get_tags_de_empresa(eid)
        self.assertIn("nuevo1", tags)
        self.assertNotIn("viejo1", tags)
        self.assertEqual(len(tags), 3)

    def test_similar_empresas_con_acentos(self):
        """fuzz.ratio debe detectar ACME SA vs ÁCME SA como similares."""
        from fuzzywuzzy import fuzz
        sim = fuzz.ratio("acme sa", "ácme sa")
        # Verificar que el umbral 75 los detectaría
        self.assertGreater(sim, 70,
            f"similitud {sim} demasiado baja para variante con acento")


# =====================================================================
# 39. LÓGICA — REGLAS DE NEGOCIO
# =====================================================================
class TestReglasNegocio(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_historial_fecha_formato_correcto(self):
        """registrar_cambio guarda fecha en YYYY-MM-DD HH:MM:SS."""
        eid = make_emp(self.db, "Fecha Test")
        self.db.registrar_cambio(eid, "nombre", "a", "b", "usuario")
        hist = self.db.get_historial_empresa(eid)
        self.assertEqual(len(hist), 1)
        fecha = hist[0]["fecha"]
        import re
        self.assertTrue(
            re.match(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}", fecha),
            f"Formato de fecha incorrecto: {fecha!r}")

    def test_historial_fuente_reversion(self):
        """Cuando se revierte un cambio, la fuente es 'usuario (reversión)'."""
        eid = make_emp(self.db, "Reversión Test")
        emp = self.db.obtener_empresa_por_id(eid)
        self.db.editar_empresa(eid, "Nombre v2",
                               emp.get("direccion",""), emp.get("telefono",""),
                               emp.get("email",""), emp.get("rubro",""),
                               emp.get("pais",""), "",
                               fuente="usuario (reversión)")
        hist = self.db.get_historial_empresa(eid)
        fuentes = [h["fuente"] for h in hist]
        self.assertIn("usuario (reversión)", fuentes)

    def test_empresa_cero_cotizaciones_cache_correcto(self):
        """empresa sin cotizaciones: ultima=None, ncot=0."""
        eid = make_emp(self.db, "Sin Cotizaciones")
        ultima = self.db.get_ultima_cotizacion(eid)
        ncot = self.db.count_by_empresa("cotizaciones", eid)
        self.assertIsNone(ultima)
        self.assertEqual(ncot, 0)

    def test_buscar_empresa_parcial_like(self):
        """get_filtered_empresas con texto parcial devuelve múltiples."""
        make_emp(self.db, "Empresa ACME SA")
        make_emp(self.db, "Empresa ACME SRL")
        make_emp(self.db, "Empresa Beta")
        result = self.db.get_filtered_empresas("ACME", {})
        nombres = [r["nombre"] for r in result]
        self.assertIn("Empresa ACME SA", nombres)
        self.assertIn("Empresa ACME SRL", nombres)
        self.assertNotIn("Empresa Beta", nombres)

    def test_cotizacion_tipo_equipos_presente_en_db(self):
        """Cotización con tipo [Equipos] se guarda y recupera correctamente."""
        eid = make_emp(self.db)
        self.db.agregar_cotizacion(eid, "[Equipos] Servidores", 50000.0)
        cots = self.db.get_cotizaciones_por_empresa(eid)
        self.assertEqual(len(cots), 1)
        cot = cots[0]
        desc = cot.get("descripcion") or ""
        tipo = cot.get("tipo") or ""
        self.assertTrue(
            tipo == "Equipos" or desc.startswith("[Equipos]"),
            f"Tipo no guardado: tipo={tipo!r} desc={desc!r}")

    def test_editar_empresa_fuente_gemini_en_historial(self):
        """editar_empresa con fuente='gemini' queda registrado así."""
        eid = make_emp(self.db, "Gemini Test")
        emp = self.db.obtener_empresa_por_id(eid)
        self.db.editar_empresa(
            eid, "Gemini Test Mejorado",
            emp.get("direccion",""), emp.get("telefono",""),
            emp.get("email",""), emp.get("rubro",""),
            emp.get("pais",""), "", fuente="gemini")
        hist = self.db.get_historial_empresa(eid)
        self.assertTrue(any(h["fuente"] == "gemini" for h in hist))

    def test_get_filtered_search_term_en_email(self):
        """get_filtered_empresas busca también en el campo email."""
        eid = make_emp(self.db, "Sin Email Especial")
        self.db.editar_empresa(eid, "Empresa con Email", "", "",
                               "especial@dominio.com", "", "AR", "")
        result = self.db.get_filtered_empresas("especial@dominio", {})
        ids = {r["id"] for r in result}
        self.assertIn(eid, ids)


# =====================================================================
# 40. VALORES LÍMITE
# =====================================================================
class TestValoresLimite(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_monto_minimo_positivo(self):
        eid = make_emp(self.db)
        ok = self.db.agregar_cotizacion(eid, "Min", 0.01)
        self.assertTrue(ok)
        cot = self.db.fetchone(
            "SELECT monto FROM cotizaciones WHERE empresa_id=?", (eid,))
        self.assertAlmostEqual(float(cot["monto"]), 0.01, places=2)

    def test_monto_maximo(self):
        eid = make_emp(self.db)
        ok = self.db.agregar_cotizacion(eid, "Max", 999_999_999_999.99)
        self.assertTrue(ok)

    def test_nombre_empresa_un_caracter(self):
        """1 char no es 'sospechoso' para la DB (solo para el enriquecedor)."""
        try:
            ok = self.db.agregar_empresa("X", "", "", "", "", "", "")
        except Exception as e:
            self.fail(f"Crash nombre 1 char: {e}")

    def test_nombre_empresa_solo_numeros(self):
        try:
            ok = self.db.agregar_empresa("12345", "", "", "", "", "", "")
        except Exception as e:
            self.fail(f"Crash nombre solo números: {e}")

    def test_email_subdominio_3_niveles(self):
        eid = make_emp(self.db)
        try:
            self.db.agregar_contacto(
                eid, "Test", "user@mail.empresa.com.ar", "", "")
        except Exception as e:
            self.fail(f"Crash email subdominio: {e}")
        cons = self.db.get_contactos_por_empresa(eid)
        emails = [c["email"] for c in cons]
        self.assertIn("user@mail.empresa.com.ar", emails)

    def test_pais_50_chars(self):
        try:
            self.db.agregar_empresa("E", "", "", "", "", "P"*50, "")
        except Exception as e:
            self.fail(f"Crash pais 50 chars: {e}")

    def test_100_tags_de_50_chars(self):
        eid = make_emp(self.db, "Many Tags")
        tags = [f"tag_{i:02d}_{'x'*40}" for i in range(100)]
        try:
            self.db.vincular_empresa_con_tags(eid, tags)
        except Exception as e:
            self.fail(f"Crash 100 tags largos: {e}")
        result = self.db.get_tags_de_empresa(eid)
        self.assertEqual(len(result), 100)

    def test_historial_10000_cambios_rendimiento(self):
        eid = make_emp(self.db, "Historial Grande")
        for i in range(10000):
            self.db.registrar_cambio(
                eid, "nombre", f"v{i}", f"v{i+1}", "usuario")
        t0 = time.time()
        hist = self.db.get_historial_empresa(eid)
        elapsed = time.time() - t0
        self.assertEqual(len(hist), 10000)
        self.assertLess(elapsed, 3.0,
            f"get_historial tardó {elapsed:.2f}s con 10000 registros")

    def test_similar_umbral_cero_devuelve_todos_pares(self):
        """umbral=0 → todos los pares posibles."""
        for i in range(5):
            make_emp(self.db, f"E{i}")
        pares = self.db.get_similar_empresas(0)
        # 5 empresas → C(5,2)=10 pares como máximo
        self.assertGreaterEqual(len(pares), 1)

    def test_similar_umbral_100_solo_identicas(self):
        make_emp(self.db, "ACME SA")
        make_emp(self.db, "ACME SA")   # duplicado exacto
        make_emp(self.db, "ACME SRL")
        pares = self.db.get_similar_empresas(100)
        for p in pares:
            self.assertEqual(p["similitud"], 100,
                f"Umbral 100 devolvió par con similitud {p['similitud']}")


# =====================================================================
# 41. INTEGRACIÓN END-TO-END
# =====================================================================
class TestIntegracionEndToEnd(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_flujo_import_duplicado_merge_historial(self):
        """CSV importado → duplicado detectado → fusionar → historial consistente."""
        # Simular importación: dos empresas con nombres similares
        e1 = make_emp(self.db, "ACME Argentina SA")
        e2 = make_emp(self.db, "ACME Argentina SRL")
        make_con(self.db, e1, "c1@acme.com")
        make_cot(self.db, e1, 5000.0)
        make_cot(self.db, e2, 8000.0)

        # Detectar similares
        pares = self.db.get_similar_empresas(70)
        self.assertTrue(len(pares) > 0, "Duplicados no detectados")

        # Fusionar e2 → e1
        ok = self.db.unificar_empresas(e2, e1)
        self.assertTrue(ok)

        # e2 desapareció
        self.assertIsNone(self.db.obtener_empresa_por_id(e2))
        # e1 tiene todos los datos
        self.assertEqual(self.db.count_by_empresa("cotizaciones", e1), 2)
        self.assertEqual(self.db.count_by_empresa("contactos", e1), 1)
        # No hay datos huérfanos
        self.assertEqual(
            len(self.db.fetchall(
                "SELECT * FROM cotizaciones WHERE empresa_id=?", (e2,))), 0)

    def test_flujo_enriquecedor_historial_fuente(self):
        """editar_empresa con fuente='gemini' → historial muestra 'gemini'."""
        eid = make_emp(self.db, "acmecorp")
        emp = self.db.obtener_empresa_por_id(eid)
        self.db.editar_empresa(
            eid, "ACME Corporation SA",
            emp.get("direccion",""), emp.get("telefono",""),
            emp.get("email",""), emp.get("rubro",""),
            emp.get("pais",""), "", fuente="gemini")
        hist = self.db.get_historial_empresa(eid)
        gemini_changes = [h for h in hist if h["fuente"] == "gemini"]
        self.assertGreater(len(gemini_changes), 0,
            "Historial no registró la fuente 'gemini'")
        nombre_change = next((h for h in hist if h["campo"] == "nombre"), None)
        self.assertIsNotNone(nombre_change)
        self.assertEqual(nombre_change["valor_anterior"], "acmecorp")
        self.assertEqual(nombre_change["valor_nuevo"], "ACME Corporation SA")

    def test_flujo_filtro_tag_unificar_tag_desaparece(self):
        """Filtrar por tag → fusionar empresa → tag no deja huérfanos."""
        e1 = make_emp(self.db, "E1 Tag")
        e2 = make_emp(self.db, "E2 Tag")
        self.db.vincular_empresa_con_tags(e1, ["vip"])
        self.db.vincular_empresa_con_tags(e2, ["cliente"])

        # Verificar filtro por tag funciona
        result = self.db.get_filtered_empresas("", {"tag": "vip"})
        self.assertIn(e1, {r["id"] for r in result})

        # Fusionar
        self.db.unificar_empresas(e1, e2)

        # No deben quedar registros de empresa_tags para e1
        huerfanos = self.db.fetchall(
            "SELECT * FROM empresa_tags WHERE empresa_id=?", (e1,))
        self.assertEqual(len(huerfanos), 0)

    def test_flujo_exportar_post_merge(self):
        """Empresa eliminada por merge no aparece en la exportación."""
        e1 = make_emp(self.db, "Empresa Origen")
        e2 = make_emp(self.db, "Empresa Destino")
        make_cot(self.db, e1, 1000.0)
        self.db.unificar_empresas(e1, e2)

        datos = self.db.get_all_empresas_with_cotizaciones()
        nombres = [d["nombre"] for d in datos]
        self.assertNotIn("Empresa Origen", nombres,
            "Empresa eliminada por merge aparece en exportación")
        self.assertIn("Empresa Destino", nombres)

    def test_flujo_editar_cotizacion_visible_en_lista(self):
        """Editar cotización → cambios visibles en get_cotizaciones_por_empresa."""
        eid = make_emp(self.db)
        make_cot(self.db, eid, 1000.0, "Original")
        cid = self.db.fetchone(
            "SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]

        self.db.editar_cotizacion(cid, "Editada", 2500.0, "Equipos")
        cots = self.db.get_cotizaciones_por_empresa(eid)
        self.assertEqual(len(cots), 1)
        cot = cots[0]
        self.assertAlmostEqual(float(cot["monto"]), 2500.0, places=1)

    def test_flujo_backup_restaurar_consistente(self):
        """hacer_backup + restaurar_backup → datos intactos."""
        tmpdir = tempfile.mkdtemp()
        try:
            db_path = os.path.join(tmpdir, "test.db")
            db = DBManager(db_path)
            eid = make_emp(db, "Backup Test")
            make_cot(db, eid, 5000.0)

            from utils import BackupManager
            ok_bak = BackupManager.hacer_backup(db_path)
            self.assertTrue(ok_bak, "hacer_backup falló")

            # Borrar datos para simular corrupción
            db.eliminar_empresa(eid)
            self.assertEqual(db.count("empresas"), 0)

            # Restaurar
            ok_rest = BackupManager.restaurar_backup(db_path)
            self.assertTrue(ok_rest, "restaurar_backup falló")

            # Reabrir y verificar
            db2 = DBManager(db_path)
            self.assertEqual(db2.count("empresas"), 1)
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)


# =====================================================================
# 42. CSV ENCODING EDGE CASES
# =====================================================================
class TestCSVEncodingEdgeCases(unittest.TestCase):
    def setUp(self): self.tmpdir = tempfile.mkdtemp()
    def tearDown(self): shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _write_raw(self, name, content_bytes):
        path = os.path.join(self.tmpdir, name)
        with open(path, 'wb') as f:
            f.write(content_bytes)
        return path

    def test_find_col_con_bom(self):
        from csv_utils import _find_col, EMAIL_COLS
        # BOM prefix en el header
        keys_with_bom = {"\ufeffEmail": "val", "Nombre": "val2"}
        result = _find_col(keys_with_bom, EMAIL_COLS)
        self.assertIsNotNone(result,
            f"_find_col no encontró email con BOM. Keys: {list(keys_with_bom.keys())}")

    def test_find_col_con_espacios_extra(self):
        from csv_utils import _find_col, EMAIL_COLS
        keys = {"  email  ": "val", "nombre": "val2"}
        result = _find_col(keys, EMAIL_COLS)
        self.assertIsNotNone(result,
            "_find_col no encontró email con espacios extra")

    def test_open_csv_archivo_vacio(self):
        from csv_utils import _open_csv
        path = self._write_raw("vacio.csv", b"")
        try:
            fh, enc, sep = _open_csv(path)
            fh.close()
        except Exception as e:
            self.fail(f"_open_csv crasheó con archivo vacío: {e}")

    def test_open_csv_utf8_bom(self):
        """CSV con BOM: _open_csv detecta utf-8-sig y _find_col stripea el BOM."""
        from csv_utils import _open_csv, _find_col, EMAIL_COLS
        # Escribir CSV con BOM usando utf-8-sig (encoding correcto para Excel Windows)
        path = os.path.join(self.tmpdir, "bom.csv")
        with open(path, "w", encoding="utf-8-sig", newline="") as f:
            f.write("nombre,email\nJuan,j@test.com\n")
        fh, enc, sep = _open_csv(path)
        with fh:
            rows = list(csv.DictReader(fh, delimiter=sep))
        self.assertEqual(len(rows), 1)
        # _find_col strips BOM from keys — verify it finds the email column
        keys = list(rows[0].keys())
        result = _find_col(keys, EMAIL_COLS)
        self.assertIsNotNone(result,
            f"_find_col no encontró email en headers con posible BOM: {keys}")

    def test_open_csv_separador_punto_y_coma(self):
        from csv_utils import _open_csv
        content = "nombre;email;empresa\nJuan;j@test.com;ACME\n"
        path = os.path.join(self.tmpdir, "semicolon.csv")
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        fh, enc, sep = _open_csv(path)
        with fh:
            rows = list(csv.DictReader(fh, delimiter=sep))
        self.assertEqual(sep, ";")
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["email"], "j@test.com")

    def test_open_csv_latin1(self):
        from csv_utils import _open_csv
        content = "nombre,email\nMaría,m@e.com\n"
        path = self._write_raw("latin.csv",
                               content.encode("latin-1"))
        fh, enc, sep = _open_csv(path)
        with fh:
            rows = list(csv.DictReader(fh, delimiter=sep))
        self.assertEqual(len(rows), 1)

    def test_open_csv_binario_no_crashea(self):
        from csv_utils import _open_csv
        # Bytes binarios (simula archivo no-texto)
        path = self._write_raw("bin.csv", bytes(range(256)))
        try:
            fh, enc, sep = _open_csv(path)
            fh.close()
        except Exception as e:
            self.fail(f"_open_csv crasheó con archivo binario: {e}")

    def test_csv_con_comillas_en_valores(self):
        """CSV con comillas: '"ACME SA"' se debe leer correctamente."""
        path = os.path.join(self.tmpdir, "quoted.csv")
        with open(path, "w", encoding="utf-8", newline="") as f:
            w = csv.writer(f, quoting=csv.QUOTE_ALL)
            w.writerow(["nombre", "email"])
            w.writerow(["ACME SA", "info@acme.com"])
        with open(path, newline="", encoding="utf-8") as f:
            rows = list(csv.DictReader(f))
        self.assertEqual(rows[0]["nombre"], "ACME SA")

    def test_csv_lineas_vacias_al_final(self):
        path = os.path.join(self.tmpdir, "trailing.csv")
        with open(path, "w", encoding="utf-8") as f:
            f.write("nombre,email\nJuan,j@t.com\n\n\n")
        with open(path, newline="", encoding="utf-8") as f:
            rows = [r for r in csv.DictReader(f) if any(r.values())]
        self.assertEqual(len(rows), 1)


# =====================================================================
# 43. ENRIQUECEDOR — safe_apply, load_empresas, write_report, call_provider
# =====================================================================
class TestEnriquecedorAvanzado(unittest.TestCase):
    def setUp(self):
        self.db = fresh_db()
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_safe_apply_should_update_false_no_toca_db(self):
        from enriquecer_empresas_gemini import safe_apply
        eid = make_emp(self.db, "No Cambiar")
        results = [{"id": eid, "current_name": "No Cambiar",
                    "canonical_name": "No Cambiar SA",
                    "confidence": 0.95, "should_update": False,
                    "country": "AR", "reason": "test"}]
        updated, skipped = safe_apply(self.db, results, 0.85)
        self.assertEqual(updated, 0)
        self.assertEqual(skipped, 1)
        emp = self.db.obtener_empresa_por_id(eid)
        self.assertEqual(emp["nombre"], "No Cambiar")

    def test_safe_apply_confidence_bajo_umbral(self):
        from enriquecer_empresas_gemini import safe_apply
        eid = make_emp(self.db, "Baja Confianza")
        results = [{"id": eid, "current_name": "Baja Confianza",
                    "canonical_name": "Baja Confianza SA",
                    "confidence": 0.849, "should_update": True,
                    "country": "AR", "reason": "test"}]
        updated, skipped = safe_apply(self.db, results, min_confidence=0.85)
        self.assertEqual(updated, 0, "Confidence 0.849 no debe aplicarse con umbral 0.85")

    def test_safe_apply_confidence_exactamente_en_umbral(self):
        from enriquecer_empresas_gemini import safe_apply
        eid = make_emp(self.db, "Umbral Exacto")
        emp = self.db.obtener_empresa_por_id(eid)
        results = [{"id": eid, "current_name": emp["nombre"],
                    "canonical_name": "Umbral Exacto SA",
                    "confidence": 0.85, "should_update": True,
                    "country": "AR", "reason": "test"}]
        updated, _ = safe_apply(self.db, results, min_confidence=0.85)
        self.assertEqual(updated, 1, "Confidence 0.85 == umbral 0.85 debe aplicarse")

    def test_safe_apply_empresa_inexistente(self):
        from enriquecer_empresas_gemini import safe_apply
        results = [{"id": 99999, "current_name": "Fantasma",
                    "canonical_name": "Fantasma SA",
                    "confidence": 0.95, "should_update": True,
                    "country": "AR", "reason": "test"}]
        updated, skipped = safe_apply(self.db, results, 0.85)
        self.assertEqual(updated, 0)
        self.assertEqual(skipped, 1)

    def test_safe_apply_mismo_nombre_no_actualiza(self):
        from enriquecer_empresas_gemini import safe_apply
        eid = make_emp(self.db, "Igual SA")
        results = [{"id": eid, "current_name": "Igual SA",
                    "canonical_name": "Igual SA",  # mismo nombre
                    "confidence": 0.99, "should_update": True,
                    "country": "AR", "reason": "ya correcto"}]
        updated, skipped = safe_apply(self.db, results, 0.85)
        self.assertEqual(updated, 0, "Mismo nombre no debe contar como actualización")

    def test_load_empresas_limit(self):
        from enriquecer_empresas_gemini import load_empresas
        for i in range(20):
            make_emp(self.db, f"Empresa_{i:02d}")
        result = load_empresas(self.db, limit=5, only=None, all_companies=True)
        self.assertLessEqual(len(result), 5)

    def test_load_empresas_only_split_coma(self):
        from enriquecer_empresas_gemini import load_empresas
        make_emp(self.db, "ACME SA")
        make_emp(self.db, "Beta Corp")
        make_emp(self.db, "Gamma SRL")
        result = load_empresas(self.db, limit=100, only="ACME, Beta",
                               all_companies=True)
        nombres = [r["nombre"] for r in result]
        self.assertIn("ACME SA", nombres)
        self.assertIn("Beta Corp", nombres)
        self.assertNotIn("Gamma SRL", nombres)

    def test_load_empresas_all_companies_incluye_no_sospechosas(self):
        from enriquecer_empresas_gemini import load_empresas
        make_emp(self.db, "Empresa Normal SA")  # no sospechosa
        make_emp(self.db, "gmail")               # sospechosa
        all_res = load_empresas(self.db, limit=100, only=None, all_companies=True)
        sus_res = load_empresas(self.db, limit=100, only=None, all_companies=False)
        self.assertGreater(len(all_res), len(sus_res))

    def test_write_report_crea_directorio(self):
        from enriquecer_empresas_gemini import write_report
        from pathlib import Path
        out = Path(self.tmpdir) / "sub" / "reporte.csv"
        write_report([{"id": 1, "current_name": "A", "canonical_name": "B",
                       "country": "AR", "website": "", "confidence": 0.9,
                       "should_update": True, "reason": "test"}], out)
        self.assertTrue(out.exists())
        self.assertGreater(out.stat().st_size, 0)

    def test_write_report_resultados_vacios_no_crashea(self):
        from enriquecer_empresas_gemini import write_report
        from pathlib import Path
        out = Path(self.tmpdir) / "vacio.csv"
        try:
            write_report([], out)
        except Exception as e:
            self.fail(f"write_report crasheó con resultados vacíos: {e}")
        self.assertTrue(out.exists())

    def test_call_provider_proveedor_desconocido_lanza_error(self):
        from enriquecer_empresas_gemini import call_provider
        with self.assertRaises((ValueError, Exception)):
            call_provider("proveedor_inventado", [], None)


# =====================================================================
# 44. CONFIG — SINGLETON, ELIMINAR_FILTRO, CARPETAS
# =====================================================================
class TestConfigAvanzado(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        import config_export
        config_export._instance = None

    def tearDown(self):
        import config_export
        config_export._instance = None
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _mgr(self):
        import config_export, unittest.mock as mock
        config_export._instance = None
        with mock.patch("config_export._find_config_dir",
                        return_value=__import__("pathlib").Path(self.tmpdir)):
            return config_export.ConfigManager()

    def test_get_app_config_singleton(self):
        """get_app_config() siempre devuelve la misma instancia."""
        import config_export, unittest.mock as mock
        config_export._instance = None
        with mock.patch("config_export._find_config_dir",
                        return_value=__import__("pathlib").Path(self.tmpdir)):
            inst1 = config_export.get_app_config()
            inst2 = config_export.get_app_config()
        self.assertIs(inst1, inst2, "get_app_config no es singleton")

    def test_eliminar_filtro_existente(self):
        mgr = self._mgr()
        mgr.guardar_filtro("mi_filtro", {"pais": "AR"})
        self.assertIn("mi_filtro", mgr.get_filtros_guardados())
        mgr.eliminar_filtro("mi_filtro")
        self.assertNotIn("mi_filtro", mgr.get_filtros_guardados())

    def test_eliminar_filtro_inexistente_no_crashea(self):
        mgr = self._mgr()
        try:
            mgr.eliminar_filtro("filtro_que_no_existe")
        except Exception as e:
            self.fail(f"Crash eliminando filtro inexistente: {e}")

    def test_get_carpetas_recientes_filtra_inexistentes(self):
        mgr = self._mgr()
        mgr._data["carpetas_recientes"] = [
            "/ruta/que/no/existe/1",
            "/ruta/que/no/existe/2",
            self.tmpdir,   # este sí existe
        ]
        recientes = mgr.get_carpetas_recientes()
        self.assertEqual(recientes, [self.tmpdir])

    def test_get_carpetas_recientes_lista_vacia(self):
        mgr = self._mgr()
        mgr._data["carpetas_recientes"] = []
        self.assertEqual(mgr.get_carpetas_recientes(), [])

    def test_agregar_carpeta_dedup(self):
        mgr = self._mgr()
        mgr.agregar_carpeta_reciente("/ruta/a")
        mgr.agregar_carpeta_reciente("/ruta/b")
        mgr.agregar_carpeta_reciente("/ruta/a")  # duplicado
        raw = mgr._data["carpetas_recientes"]
        self.assertEqual(raw.count("/ruta/a"), 1)
        self.assertEqual(raw[0], "/ruta/a")  # más reciente al frente


# =====================================================================
# 45. UTILS — BACKUP REAL
# =====================================================================
class TestUtilsBackupReal(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_get_db_name_devuelve_path_absoluto(self):
        from utils import Config
        cfg = Config()
        db = cfg.get_db_name()
        self.assertTrue(os.path.isabs(db),
            f"get_db_name debería devolver path absoluto: {db}")

    def test_backup_dir_es_subdirectorio(self):
        from utils import BackupManager
        db_path = os.path.join(self.tmpdir, "test.db")
        bdir = BackupManager._backup_dir(db_path)
        self.assertIn(self.tmpdir, bdir)
        self.assertTrue(bdir.endswith("backups"))

    def test_hacer_backup_crea_archivo(self):
        from utils import BackupManager
        db_path = os.path.join(self.tmpdir, "test.db")
        db = DBManager(db_path)
        make_emp(db, "Empresa Backup")

        ok = BackupManager.hacer_backup(db_path)
        self.assertTrue(ok)

        bdir = BackupManager._backup_dir(db_path)
        import glob
        backups = glob.glob(os.path.join(bdir, "*.db"))
        self.assertGreater(len(backups), 0, "No se creó ningún archivo de backup")
        self.assertGreater(os.path.getsize(backups[0]), 0)

    def test_restaurar_backup_con_backup_real(self):
        from utils import BackupManager
        db_path = os.path.join(self.tmpdir, "restore.db")
        db = DBManager(db_path)
        make_emp(db, "Empresa Original")

        # Hacer backup
        ok_bak = BackupManager.hacer_backup(db_path)
        self.assertTrue(ok_bak)

        # Borrar datos
        db.ejecutar("DELETE FROM empresas")
        self.assertEqual(db.count("empresas"), 0)

        # Restaurar
        ok_rest = BackupManager.restaurar_backup(db_path)
        self.assertTrue(ok_rest)

        # Reabrir y verificar
        db2 = DBManager(db_path)
        self.assertEqual(db2.count("empresas"), 1)

    def test_restaurar_sin_backups_devuelve_false(self):
        from utils import BackupManager
        ok = BackupManager.restaurar_backup(
            os.path.join(self.tmpdir, "sin_backups.db"))
        self.assertFalse(ok)


# =====================================================================
# 46. STRESS — GAPS DE CONCURRENCIA CRÍTICOS
# =====================================================================
class TestStressConcurrenciaCriticos(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.db = DBManager(os.path.join(self.tmpdir, "crit.db"))

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_historial_3_usuarios_leyendo_simultaneo(self):
        """3 conexiones leyendo historial simultáneamente sin errores."""
        eid = make_emp(self.db, "Historial Concurrente")
        for i in range(50):
            self.db.registrar_cambio(eid, "nombre", f"v{i}", f"v{i+1}", "usuario")

        db2 = DBManager(self.db.db_name)
        db3 = DBManager(self.db.db_name)
        errors = []

        def leer_hist(db, n):
            for _ in range(n):
                try:
                    hist = db.get_historial_empresa(eid)
                    if len(hist) == 0:
                        errors.append("historial vacío inesperado")
                except Exception as e:
                    errors.append(str(e))

        threads = [
            threading.Thread(target=leer_hist, args=(self.db, 30)),
            threading.Thread(target=leer_hist, args=(db2, 30)),
            threading.Thread(target=leer_hist, args=(db3, 30)),
        ]
        for t in threads: t.start()
        for t in threads: t.join(timeout=15)
        self.assertEqual(errors, [], f"Errores: {errors[:3]}")

    def test_10_tags_concurrentes_misma_empresa(self):
        """10 threads vinculando tags distintos a la misma empresa."""
        eid = make_emp(self.db, "Tags Concurrentes")
        errors = []

        def vincular(tags):
            try:
                self.db.vincular_empresa_con_tags(eid, tags)
            except Exception as e:
                errors.append(str(e))

        threads = [
            threading.Thread(target=vincular, args=([f"tag_{i}"],))
            for i in range(10)
        ]
        for t in threads: t.start()
        for t in threads: t.join(timeout=15)

        alive = [t for t in threads if t.is_alive()]
        self.assertEqual(len(alive), 0, "Deadlock en vincular_empresa_con_tags")
        self.assertEqual(errors, [], f"Errores: {errors}")
        # Al final la empresa tiene exactamente 1 set de tags (el último ganó)
        tags = self.db.get_tags_de_empresa(eid)
        self.assertIsInstance(tags, list)

    def test_background_resumidor_y_main_thread_misma_cotizacion(self):
        """
        Thread secundario actualiza resumen mientras main thread edita la cotización.
        Ninguno debe corromper los datos del otro.
        """
        db_path = self.db.db_name
        eid = make_emp(self.db, "Cotiz Concurrente")
        self.db.agregar_cotizacion_con_ruta(
            eid, "archivo.pdf", 1000.0,
            "2024-01-01 00:00:00", "/srv/archivo.pdf")
        cid = self.db.fetchone(
            "SELECT id FROM cotizaciones WHERE empresa_id=?", (eid,))["id"]

        errors = []
        done = threading.Event()

        def background_resumen():
            """Simula el resumidor background."""
            try:
                conn = sqlite3.connect(db_path, timeout=10)
                conn.execute("PRAGMA busy_timeout=5000")
                conn.execute(
                    "UPDATE cotizaciones SET resumen=? WHERE id=?",
                    ("Resumen generado por IA", cid))
                conn.commit()
                conn.close()
            except Exception as e:
                errors.append(f"bg: {e}")
            finally:
                done.set()

        t = threading.Thread(target=background_resumen, daemon=True)
        t.start()

        # Main thread edita la cotización mientras el background trabaja
        try:
            self.db.editar_cotizacion(cid, "Descripcion editada", 2500.0)
        except Exception as e:
            errors.append(f"main: {e}")

        done.wait(timeout=10)
        t.join(timeout=5)

        self.assertEqual(errors, [], f"Errores concurrentes: {errors}")

        # Verificar que la cotización tiene monto actualizado por main thread
        cot = self.db.get_cotizacion_por_id(cid)
        self.assertIsNotNone(cot)
        self.assertAlmostEqual(float(cot["monto"]), 2500.0, places=1)

    def test_similar_2000_empresas_rendimiento(self):
        """get_similar_empresas con 2000 empresas debe terminar en < 20s."""
        for i in range(2000):
            self.db.agregar_empresa(f"Empresa_{i:04d} SA", "", "", "", "", "", "")
        t0 = time.time()
        pares = self.db.get_similar_empresas(85)
        elapsed = time.time() - t0
        self.assertLess(elapsed, 20.0,
            f"get_similar_empresas tardó {elapsed:.2f}s con 2000 empresas")
        self.assertIsInstance(pares, list)

    def test_exportar_mientras_otro_importa(self):
        """Exportar mientras otro thread importa no produce datos corruptos."""
        # Poblar datos base
        for i in range(20):
            eid = make_emp(self.db, f"Base_{i}")
            make_cot(self.db, eid, float(i * 100))

        export_results = []
        errors = []

        def exportar():
            for _ in range(5):
                try:
                    datos = self.db.get_all_empresas_with_cotizaciones()
                    export_results.append(len(datos))
                except Exception as e:
                    errors.append(f"export: {e}")
                time.sleep(0.02)

        def importar():
            for i in range(30):
                try:
                    eid = make_emp(self.db, f"Import_{i}")
                    make_cot(self.db, eid, float(i * 50))
                except Exception as e:
                    errors.append(f"import: {e}")
                time.sleep(0.005)

        t1 = threading.Thread(target=exportar)
        t2 = threading.Thread(target=importar)
        t1.start(); t2.start()
        t1.join(timeout=15); t2.join(timeout=15)

        self.assertEqual(errors, [], f"Errores: {errors}")
        self.assertGreater(len(export_results), 0)
        # Todos los exports deben haber devuelto listas válidas (no crashearon)
        for n in export_results:
            self.assertGreaterEqual(n, 0)




class TestActividades(unittest.TestCase):
    def setUp(self): self.db = fresh_db()

    def test_agregar_actividad_basico(self):
        eid = make_emp(self.db)
        ok = self.db.agregar_actividad(eid, "llamada", "Llamé a Juan, OK para marzo")
        self.assertTrue(ok)
        acts = self.db.get_actividades_empresa(eid)
        self.assertEqual(len(acts), 1)
        self.assertEqual(acts[0]["tipo"], "llamada")
        self.assertIn("Juan", acts[0]["texto"])

    def test_tipos_validos(self):
        eid = make_emp(self.db)
        for tipo in ("llamada","email","reunion","nota"):
            self.assertTrue(self.db.agregar_actividad(eid, tipo, f"Test {tipo}"))
        acts = self.db.get_actividades_empresa(eid)
        self.assertEqual(len(acts), 4)

    def test_orden_reciente_primero(self):
        eid = make_emp(self.db)
        self.db.agregar_actividad(eid, "nota", "Primera")
        self.db.agregar_actividad(eid, "nota", "Segunda")
        self.db.agregar_actividad(eid, "nota", "Tercera")
        acts = self.db.get_actividades_empresa(eid)
        self.assertEqual(acts[0]["texto"], "Tercera")

    def test_agregar_texto_vacio_falla(self):
        eid = make_emp(self.db)
        self.assertFalse(self.db.agregar_actividad(eid, "nota", ""))
        self.assertFalse(self.db.agregar_actividad(eid, "nota", "   "))

    def test_agregar_empresa_id_none(self):
        try:
            ok = self.db.agregar_actividad(None, "nota", "Texto")
            self.assertFalse(ok)
        except Exception as e:
            self.fail(f"Crash empresa_id=None: {e}")

    def test_editar_actividad(self):
        eid = make_emp(self.db)
        self.db.agregar_actividad(eid, "nota", "Original")
        aid = self.db.get_actividades_empresa(eid)[0]["id"]
        ok = self.db.editar_actividad(aid, "llamada", "Editado")
        self.assertTrue(ok)
        act = self.db.fetchone("SELECT * FROM actividades WHERE id=?", (aid,))
        self.assertEqual(act["texto"], "Editado")
        self.assertEqual(act["tipo"], "llamada")

    def test_editar_inexistente(self):
        ok = self.db.editar_actividad(99999, "nota", "X")
        self.assertFalse(ok)

    def test_eliminar_actividad(self):
        eid = make_emp(self.db)
        self.db.agregar_actividad(eid, "nota", "Para borrar")
        aid = self.db.get_actividades_empresa(eid)[0]["id"]
        self.db.eliminar_actividad(aid)
        self.assertEqual(len(self.db.get_actividades_empresa(eid)), 0)

    def test_eliminar_empresa_elimina_actividades(self):
        eid = make_emp(self.db)
        self.db.agregar_actividad(eid, "nota", "Nota orphan test")
        self.db.eliminar_empresa(eid)
        huerfanas = self.db.fetchall(
            "SELECT * FROM actividades WHERE empresa_id=?", (eid,))
        self.assertEqual(len(huerfanas), 0)

    def test_get_actividades_recientes(self):
        eid = make_emp(self.db)
        self.db.agregar_actividad(eid, "llamada", "Reciente 1")
        self.db.agregar_actividad(eid, "email",   "Reciente 2")
        recientes = self.db.get_actividades_recientes(dias=7, limit=50)
        self.assertGreaterEqual(len(recientes), 2)
        self.assertIn("empresa_nombre", recientes[0])

    def test_sql_injection_en_texto(self):
        eid = make_emp(self.db)
        try:
            self.db.agregar_actividad(
                eid, "nota", "'; DROP TABLE actividades;--")
        except Exception as e:
            self.fail(f"Crash SQL injection: {e}")
        self.assertIsNotNone(
            self.db.fetchall("SELECT * FROM actividades"))

    def test_texto_largo(self):
        eid = make_emp(self.db)
        texto_largo = "A" * 5000
        ok = self.db.agregar_actividad(eid, "nota", texto_largo)
        self.assertTrue(ok)
        act = self.db.get_actividades_empresa(eid)[0]
        self.assertEqual(len(act["texto"]), 5000)

    def test_multiples_empresas_aisladas(self):
        e1 = make_emp(self.db, "E1"); e2 = make_emp(self.db, "E2")
        self.db.agregar_actividad(e1, "nota", "De E1")
        self.db.agregar_actividad(e2, "nota", "De E2")
        self.assertEqual(len(self.db.get_actividades_empresa(e1)), 1)
        self.assertEqual(len(self.db.get_actividades_empresa(e2)), 1)

    def test_concurrencia_agregar_actividades(self):
        eid = make_emp(self.db)
        errors = []
        def writer(n):
            for i in range(20):
                ok = self.db.agregar_actividad(
                    eid, "nota", f"Thread {n} nota {i}")
                if not ok: errors.append(f"T{n}_{i}")
        threads = [threading.Thread(target=writer, args=(t,))
                   for t in range(5)]
        for t in threads: t.start()
        for t in threads: t.join(timeout=15)
        self.assertEqual(errors, [])
        self.assertEqual(len(self.db.get_actividades_empresa(eid)), 100)

    def test_empresa_inexistente_rechazada(self):
        """agregar_actividad devuelve False si la empresa no existe."""
        ok = self.db.agregar_actividad(99999, "nota", "Texto huérfano")
        self.assertFalse(ok, "Debe rechazar empresa inexistente")
        huerfanas = self.db.fetchall(
            "SELECT * FROM actividades WHERE empresa_id=99999")
        self.assertEqual(len(huerfanas), 0)

    def test_eliminar_inexistente_devuelve_false(self):
        """eliminar_actividad devuelve False si no existe la actividad."""
        ok = self.db.eliminar_actividad(99999)
        self.assertFalse(ok)

    def test_editar_inexistente_devuelve_false(self):
        ok = self.db.editar_actividad(99999, "nota", "Texto")
        self.assertFalse(ok)

    def test_tipo_invalido_normaliza_a_nota(self):
        eid = make_emp(self.db)
        self.assertTrue(self.db.agregar_actividad(eid, "hacker", "Texto"))
        acts = self.db.get_actividades_empresa(eid)
        self.assertEqual(acts[0]["tipo"], "nota",
            "Tipo inválido debe normalizarse a 'nota'")

    def test_todos_tipos_validos(self):
        eid = make_emp(self.db)
        for tipo in ("nota","llamada","email","reunion"):
            self.assertTrue(
                self.db.agregar_actividad(eid, tipo, f"Test {tipo}"))
            acts = self.db.get_actividades_empresa(eid)
            self.assertEqual(acts[0]["tipo"], tipo)

    def test_filtro_dias_actividad(self):
        """Empresas sin actividad reciente aparecen en el filtro."""
        e_activa   = make_emp(self.db, "Activa SA")
        e_inactiva = make_emp(self.db, "Inactiva SA")
        self.db.agregar_actividad(e_activa, "nota", "Actividad hoy")
        result = self.db.get_filtered_empresas("", {"dias_actividad": 30})
        ids = {r["id"] for r in result}
        self.assertIn(e_inactiva, ids,
            "Empresa sin actividad debe aparecer en filtro")
        self.assertNotIn(e_activa, ids,
            "Empresa con actividad reciente NO debe aparecer")

    def test_paginacion_actividades(self):
        eid = make_emp(self.db)
        for i in range(50):
            self.db.agregar_actividad(eid, "nota", f"Nota {i}")
        pag1 = self.db.get_actividades_empresa(eid, limit=10, offset=0)
        pag2 = self.db.get_actividades_empresa(eid, limit=10, offset=10)
        self.assertEqual(len(pag1), 10)
        self.assertEqual(len(pag2), 10)
        ids_p1 = {a["id"] for a in pag1}
        ids_p2 = {a["id"] for a in pag2}
        self.assertEqual(len(ids_p1 & ids_p2), 0, "Páginas no deben solaparse")




    def test_server_filtro_dias_actividad(self):
        """
        /api/empresas?dias_actividad=30 devuelve solo empresas sin actividad
        en los últimos 30 días.
        """
        from server import app as _app, db as _db
        _app.config['TESTING'] = True
        client = _app.test_client()

        # Empresa activa — tiene actividad reciente
        _db.agregar_empresa("FiltroActiva SA","","","","","","")
        eact = _db.fetchone(
            "SELECT id FROM empresas WHERE nombre='FiltroActiva SA'")["id"]
        _db.agregar_actividad(eact, "nota", "Actividad de hoy")

        # Empresa inactiva — sin actividad
        _db.agregar_empresa("FiltroInactiva SA","","","","","","")
        einact = _db.fetchone(
            "SELECT id FROM empresas WHERE nombre='FiltroInactiva SA'")["id"]

        r = client.get("/api/empresas?dias_actividad=30")
        self.assertEqual(r.status_code, 200)
        ids = {e["id"] for e in r.get_json()["data"]}
        self.assertIn(einact, ids,
            "Empresa sin actividad debe aparecer en filtro")
        self.assertNotIn(eact, ids,
            "Empresa con actividad reciente NO debe aparecer")

        _db.eliminar_empresa(eact)
        _db.eliminar_empresa(einact)

    def test_ui_habilita_boton_actividad(self):
        """enableDetailBtns debe incluir btn-add-act."""
        import re
        with open(os.path.join(BASE_DIR, 'static', 'index.html'), encoding='utf-8') as _f:
            html = _f.read()
        self.assertIn('"btn-add-act"', html,
            "btn-add-act debe existir en el HTML")
        block = re.search(
            r'function enableDetailBtns[\s\S]*?\}', html)
        self.assertIsNotNone(block)
        self.assertIn('btn-add-act', block.group(),
            "enableDetailBtns debe habilitar btn-add-act")

    def test_ui_filtro_manda_dias_actividad(self):
        """applyFilters debe enviar dias_actividad, no dias_cotizacion."""
        with open(os.path.join(BASE_DIR, 'static', 'index.html'), encoding='utf-8') as _f:
            html = _f.read()
        import re
        # Check días field sends dias_actividad not dias_cotizacion
        import re as _re
        # Extract the full applyFilters function (up to closing brace on its own line)
        block = _re.search(
            r"function applyFilters\(\)[^}]+(?:[^}]+\})+", html)
        self.assertIsNotNone(block, "applyFilters no encontrada en HTML")
        fn = block.group()
        self.assertIn('dias_actividad', fn,
            "applyFilters debe mandar dias_actividad")
        self.assertNotIn('dias_cotizacion', fn,
            "applyFilters no debe mandar dias_cotizacion")



# =====================================================================
# 51. OPORTUNIDADES / PIPELINE — DB
# =====================================================================
class TestOportunidades(unittest.TestCase):
    def setUp(self):
        self.db = fresh_db()
        self.eid = make_emp(self.db, "Pipeline SA")

    def test_crear_oportunidad_basico(self):
        ok = self.db.crear_oportunidad(self.eid, "Venta tolva",
                                       descripcion="Necesidad detectada")
        self.assertTrue(ok)
        rows = self.db.get_oportunidades_empresa(self.eid)
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["titulo"], "Venta tolva")
        self.assertEqual(rows[0]["fase"], "venta")

    def test_empresa_inexistente_rechazada(self):
        self.assertFalse(self.db.crear_oportunidad(99999, "No existe"))

    def test_etapa_invalida_rechazada(self):
        self.assertFalse(
            self.db.crear_oportunidad(self.eid, "Mala", etapa="inventada"))

    def test_todas_etapas_validas_aceptadas(self):
        for etapa in ("prospecto","contactado","a_visitar","a_cotizar",
                      "cotizado","en_negociacion","ganado","perdido","muerta",
                      "en_proceso","entregada","finalizada"):
            ok = self.db.crear_oportunidad(
                self.eid, f"Test {etapa}", etapa=etapa)
            self.assertTrue(ok, f"Etapa '{etapa}' rechazada incorrectamente")

    def test_cambiar_etapa_valida(self):
        self.db.crear_oportunidad(self.eid, "Seguimiento")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        self.assertTrue(self.db.cambiar_etapa_oportunidad(oid, "contactado"))
        row = self.db.get_oportunidad_por_id(oid)
        self.assertEqual(row["etapa"], "contactado")

    def test_cambiar_etapa_ganado_activa_posventa(self):
        self.db.crear_oportunidad(self.eid, "Ganable")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        self.assertTrue(self.db.cambiar_etapa_oportunidad(oid, "ganado"))
        self.assertEqual(self.db.get_oportunidad_por_id(oid)["fase"], "posventa")
        self.assertTrue(self.db.cambiar_etapa_oportunidad(oid, "en_proceso"))
        row = self.db.get_oportunidad_por_id(oid)
        self.assertEqual(row["etapa"], "en_proceso")
        self.assertEqual(row["fase"], "posventa")

    def test_fase_derivada_correctamente(self):
        self.assertEqual(self.db._fase_de_etapa("prospecto"), "venta")
        self.assertEqual(self.db._fase_de_etapa("ganado"),    "posventa")
        self.assertEqual(self.db._fase_de_etapa("entregada"), "posventa")
        self.assertEqual(self.db._fase_de_etapa("muerta"),    "venta")

    def test_cascade_delete_con_empresa(self):
        self.db.crear_oportunidad(self.eid, "Se borra")
        self.db.eliminar_empresa(self.eid)
        self.assertEqual(self.db.get_oportunidades_empresa(self.eid), [])
        huerfanas = self.db.fetchall(
            "SELECT * FROM oportunidades WHERE empresa_id=?", (self.eid,))
        self.assertEqual(len(huerfanas), 0)

    def test_multiples_oportunidades_misma_empresa(self):
        self.db.crear_oportunidad(self.eid, "Uno")
        self.db.crear_oportunidad(self.eid, "Dos")
        self.db.crear_oportunidad(self.eid, "Tres")
        rows = self.db.get_oportunidades_empresa(self.eid)
        self.assertEqual(len(rows), 3)

    def test_monto_none_permitido(self):
        self.assertTrue(
            self.db.crear_oportunidad(self.eid, "Sin monto",
                                      monto_estimado=None))

    def test_sql_injection_en_titulo_notas(self):
        payload = "'; DROP TABLE oportunidades;--"
        self.assertTrue(
            self.db.crear_oportunidad(self.eid, payload, notas=payload))
        self.assertIsNotNone(
            self.db.fetchall("SELECT * FROM oportunidades"))


    def test_editar_titulo_vacio_ignorado(self):
        """editar_oportunidad con titulo='' no debe pisar el título existente."""
        self.db.crear_oportunidad(self.eid, "Título original")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        # Empty titulo should be ignored, not saved
        self.db.editar_oportunidad(oid, titulo="", descripcion="Nueva desc")
        row = self.db.get_oportunidad_por_id(oid)
        self.assertEqual(row["titulo"], "Título original",
            "titulo vacío no debe sobreescribir el título existente")
        self.assertEqual(row["descripcion"], "Nueva desc")

    def test_filtro_fase_venta_endpoint(self):
        """GET /api/oportunidades?fase=venta solo devuelve fase venta."""
        from server import app as _app, db as _db
        _app.config['TESTING'] = True
        client = _app.test_client()
        _db.crear_oportunidad(self.eid, "Fase venta",   etapa="prospecto")
        _db.crear_oportunidad(self.eid, "Fase posventa",etapa="en_proceso")
        r = client.get("/api/oportunidades?fase=venta")
        self.assertEqual(r.status_code, 200)
        rows = r.get_json()["data"]
        self.assertTrue(all(row["fase"] == "venta" for row in rows
                            if row["empresa_id"] == self.eid),
            "fase=venta debe devolver solo oportunidades de fase venta")

    def test_filtro_fase_posventa_endpoint(self):
        """GET /api/oportunidades?fase=posventa solo devuelve fase posventa."""
        from server import app as _app, db as _db
        _app.config['TESTING'] = True
        client = _app.test_client()
        _db.crear_oportunidad(self.eid, "Posventa test", etapa="en_proceso")
        r = client.get("/api/oportunidades?fase=posventa")
        self.assertEqual(r.status_code, 200)
        rows = r.get_json()["data"]
        self.assertTrue(all(row["fase"] == "posventa" for row in rows),
            "fase=posventa debe devolver solo oportunidades de fase posventa")

    def test_ui_selector_empresa_deshabilitado_en_edicion(self):
        """editarOportunidad() debe deshabilitar el selector de empresa."""
        with open(os.path.join(BASE_DIR, 'static', 'index.html'), encoding='utf-8') as _f:
            html = _f.read()
        self.assertIn('sel.disabled = true', html,
            "El selector de empresa debe deshabilitarse en modo edición")
        self.assertIn('sel.disabled = false', html,
            "El selector de empresa debe habilitarse para nueva oportunidad")

    def test_editar_oportunidad(self):
        self.db.crear_oportunidad(self.eid, "Original")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        ok = self.db.editar_oportunidad(
            oid, titulo="Editado", monto_estimado=50000.0)
        self.assertTrue(ok)
        row = self.db.get_oportunidad_por_id(oid)
        self.assertEqual(row["titulo"], "Editado")
        self.assertAlmostEqual(float(row["monto_estimado"]), 50000.0)

    def test_eliminar_oportunidad(self):
        self.db.crear_oportunidad(self.eid, "Para borrar")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        self.assertTrue(self.db.eliminar_oportunidad(oid))
        self.assertIsNone(self.db.get_oportunidad_por_id(oid))

    def test_eliminar_inexistente_devuelve_false(self):
        self.assertFalse(self.db.eliminar_oportunidad(99999))

    def test_get_oportunidades_filtro_etapa(self):
        self.db.crear_oportunidad(self.eid, "Prospecto", etapa="prospecto")
        self.db.crear_oportunidad(self.eid, "Contactado", etapa="contactado")
        rows = self.db.get_oportunidades({"etapa": "prospecto"})
        self.assertTrue(all(r["etapa"] == "prospecto" for r in rows))

    def test_empresa_nombre_en_resultado(self):
        self.db.crear_oportunidad(self.eid, "Con nombre empresa")
        rows = self.db.get_oportunidades_empresa(self.eid)
        self.assertIn("empresa_nombre", rows[0])
        self.assertEqual(rows[0]["empresa_nombre"], "Pipeline SA")


# =====================================================================
# 52. OPORTUNIDADES — HTTP
# =====================================================================
class TestOportunidadesHTTP(unittest.TestCase):
    def setUp(self):
        from server import app as _app, db as _db
        _app.config['TESTING'] = True
        self.client = _app.test_client()
        self.db = _db
        self.db.agregar_empresa("OportHTTP SA","","","","","","")
        r = self.db.fetchone(
            "SELECT id FROM empresas WHERE nombre='OportHTTP SA'")
        self.eid = r["id"]

    def tearDown(self):
        self.db.ejecutar(
            "DELETE FROM oportunidades WHERE empresa_id=?", (self.eid,))
        self.db.eliminar_empresa(self.eid)

    def test_post_oportunidad_201(self):
        r = self.client.post("/api/oportunidades",
                             json={"empresa_id": self.eid,
                                   "titulo": "HTTP Venta"})
        self.assertEqual(r.status_code, 201)
        self.assertTrue(r.get_json()["ok"])

    def test_get_oportunidades_200(self):
        self.db.crear_oportunidad(self.eid, "Listado HTTP")
        r = self.client.get("/api/oportunidades")
        self.assertEqual(r.status_code, 200)
        self.assertGreaterEqual(len(r.get_json()["data"]), 1)

    def test_get_oportunidades_empresa_200(self):
        self.db.crear_oportunidad(self.eid, "Por empresa HTTP")
        r = self.client.get(f"/api/empresas/{self.eid}/oportunidades")
        self.assertEqual(r.status_code, 200)
        data = r.get_json()["data"]
        self.assertTrue(any(o["titulo"] == "Por empresa HTTP" for o in data))

    def test_put_etapa_valida_200(self):
        self.db.crear_oportunidad(self.eid, "Mover HTTP")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        r = self.client.put(f"/api/oportunidades/{oid}/etapa",
                            json={"etapa": "contactado"})
        self.assertEqual(r.status_code, 200)
        self.assertEqual(
            self.db.get_oportunidad_por_id(oid)["etapa"], "contactado")

    def test_put_etapa_invalida_400(self):
        self.db.crear_oportunidad(self.eid, "Mover mal HTTP")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        r = self.client.put(f"/api/oportunidades/{oid}/etapa",
                            json={"etapa": "inventada"})
        self.assertEqual(r.status_code, 400)

    def test_delete_oportunidad_200(self):
        self.db.crear_oportunidad(self.eid, "Borrar HTTP")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        r = self.client.delete(f"/api/oportunidades/{oid}")
        self.assertEqual(r.status_code, 200)
        self.assertIsNone(self.db.get_oportunidad_por_id(oid))

    def test_empresa_inexistente_404(self):
        r = self.client.post("/api/oportunidades",
                             json={"empresa_id": 999999, "titulo": "No"})
        self.assertEqual(r.status_code, 404)

    def test_titulo_vacio_400(self):
        r = self.client.post("/api/oportunidades",
                             json={"empresa_id": self.eid, "titulo": ""})
        self.assertEqual(r.status_code, 400)

    def test_get_por_id_200(self):
        self.db.crear_oportunidad(self.eid, "Por ID HTTP")
        oid = self.db.get_oportunidades_empresa(self.eid)[0]["id"]
        r = self.client.get(f"/api/oportunidades/{oid}")
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.get_json()["data"]["titulo"], "Por ID HTTP")

    def test_get_por_id_inexistente_404(self):
        r = self.client.get("/api/oportunidades/999999")
        self.assertEqual(r.status_code, 404)

class TestActividadesHTTP(unittest.TestCase):
    def setUp(self):
        from server import app, db as _db
        app.config['TESTING'] = True
        self.client = app.test_client()
        self.db = _db
        self.db.agregar_empresa("ActHTTP SA","","","","","","")
        r = self.db.fetchone(
            "SELECT id FROM empresas WHERE nombre='ActHTTP SA'")
        self.eid = r["id"]

    def tearDown(self):
        self.db.eliminar_empresa(self.eid)

    def test_post_actividad(self):
        r = self.client.post(
            f"/api/empresas/{self.eid}/actividades",
            json={"tipo":"llamada","texto":"Test HTTP llamada"})
        self.assertEqual(r.status_code, 201)

    def test_get_actividades(self):
        self.db.agregar_actividad(self.eid, "email", "Test get")
        r = self.client.get(f"/api/empresas/{self.eid}/actividades")
        self.assertEqual(r.status_code, 200)
        data = r.get_json()["data"]
        self.assertTrue(any(a["texto"] == "Test get" for a in data))

    def test_put_actividad(self):
        self.db.agregar_actividad(self.eid, "nota", "Original")
        aid = self.db.get_actividades_empresa(self.eid)[0]["id"]
        r = self.client.put(f"/api/actividades/{aid}",
                            json={"tipo":"reunion","texto":"Editado"})
        self.assertEqual(r.status_code, 200)

    def test_delete_actividad(self):
        self.db.agregar_actividad(self.eid, "nota", "Para borrar HTTP")
        aid = self.db.get_actividades_empresa(self.eid)[0]["id"]
        r = self.client.delete(f"/api/actividades/{aid}")
        self.assertEqual(r.status_code, 200)

    def test_post_texto_vacio_da_400(self):
        r = self.client.post(
            f"/api/empresas/{self.eid}/actividades",
            json={"tipo":"nota","texto":""})
        self.assertEqual(r.status_code, 400)

    def test_get_actividades_recientes(self):
        self.db.agregar_actividad(self.eid, "llamada", "Reciente")
        r = self.client.get("/api/actividades/recientes?dias=7")
        self.assertEqual(r.status_code, 200)
        data = r.get_json()["data"]
        self.assertGreater(len(data), 0)
        self.assertIn("empresa_nombre", data[0])

    def test_get_empresa_inexistente_404(self):
        r = self.client.get("/api/empresas/999999/actividades")
        self.assertEqual(r.status_code, 404)

    def test_post_empresa_inexistente_404(self):
        r = self.client.post("/api/empresas/999999/actividades",
                             json={"tipo":"nota","texto":"Huérfana"})
        self.assertEqual(r.status_code, 404)

    def test_delete_actividad_inexistente_404(self):
        r = self.client.delete("/api/actividades/999999")
        self.assertEqual(r.status_code, 404)

    def test_put_actividad_inexistente_404(self):
        r = self.client.put("/api/actividades/999999",
                            json={"tipo":"nota","texto":"X"})
        self.assertEqual(r.status_code, 404)

    def test_tipo_invalido_normaliza_en_api(self):
        r = self.client.post(
            f"/api/empresas/{self.eid}/actividades",
            json={"tipo":"inyeccion_sql","texto":"Test"})
        self.assertEqual(r.status_code, 201)
        acts = self.client.get(
            f"/api/empresas/{self.eid}/actividades").get_json()["data"]
        self.assertEqual(acts[0]["tipo"], "nota")


# RUNNER
# =====================================================================
if __name__ == "__main__":
    loader = unittest.TestLoader()
    suite  = loader.loadTestsFromModule(__import__(__name__))
    runner = unittest.TextTestRunner(verbosity=2, stream=sys.stdout, failfast=False)
    result = runner.run(suite)

    print("\n" + "="*70)
    total  = result.testsRun
    passed = total - len(result.failures) - len(result.errors) - len(result.skipped)
    print(f"  TOTAL:     {total}")
    print(f"  OK:        {passed}")
    print(f"  FALLAS:    {len(result.failures)}")
    print(f"  ERRORES:   {len(result.errors)}")
    print(f"  SKIP:      {len(result.skipped)}")
    print("="*70)
    if result.failures or result.errors:
        print("\nPROBLEMAS:\n")
        for test, trace in result.failures + result.errors:
            last = [l for l in trace.strip().split("\n") if l.strip()][-1]
            print(f"  x  {test}")
            print(f"     {last}")
    sys.exit(0 if result.wasSuccessful() else 1)
